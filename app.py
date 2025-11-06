from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_from_directory, make_response, send_file
import psycopg
from psycopg.rows import dict_row
from psycopg.errors import InvalidCatalogName
import os
from dotenv import load_dotenv, find_dotenv
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, timedelta
import secrets
import re
from functools import wraps
import time
import random
import io
import sys
import smtplib
from email.message import EmailMessage
import mimetypes
import json
import unicodedata

# Carrega .env de forma robusta (procura subindo diretórios)
try:
    _dotenv_path = find_dotenv(usecwd=True)
    if _dotenv_path:
        load_dotenv(_dotenv_path)
        print(f"[ENV] .env carregado de: {_dotenv_path}")
    else:
        # Fallback: tenta .env ao lado do app.py
        _local_env = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
        if os.path.exists(_local_env):
            load_dotenv(_local_env)
            print(f"[ENV] .env carregado de: {_local_env}")
        else:
            # Tenta .env na pasta pai
            _parent_env = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '.env')
            if os.path.exists(_parent_env):
                load_dotenv(_parent_env)
                print(f"[ENV] .env carregado de: {_parent_env}")
except Exception as _e:
    print(f"[ENV] Falha ao carregar .env: {_e}")

app = Flask(__name__, static_folder='static', template_folder='templates')
# Tipos MIME explícitos para Office (garantem Content-Type correto em assets estáticos)
mimetypes.add_type('application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.docx')
mimetypes.add_type('application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', '.xlsx')
mimetypes.add_type('application/vnd.ms-excel', '.xls')
app.secret_key = os.getenv('SECRET_KEY', 'inprolib_secret_key_2024')
# Sessões permanentes quando "Lembrar-me" marcado: 30 dias
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=30)
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload
ADMIN_SETUP_TOKEN = os.getenv('ADMIN_SETUP_TOKEN', 'setup_admin_2024')
ADMIN_TEMP_PASSWORD = os.getenv('ADMIN_TEMP_PASSWORD', 'Adm@2025!')
# Expiração do token de recuperação em segundos (padrão: 60 segundos)
RESET_TOKEN_EXP_SECONDS = int(os.getenv('RESET_TOKEN_EXP_SECONDS', '60'))

# Configuração do banco de dados PostgreSQL (via variáveis de ambiente)
DB_CONFIG = {
    'dbname': os.getenv('DB_NAME', 'inprolib_schema'),
    'user': os.getenv('DB_USER', 'postgres'),
    'password': os.getenv('DB_PASSWORD', 'postgres'),
    'host': os.getenv('DB_HOST', 'localhost'),
    'port': os.getenv('DB_PORT', '5432')
}

# Garantir que a pasta de uploads exista
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs'), exist_ok=True)

# Rate limiting simples em memória: chave por IP+rota
RATE_LIMIT = {}

# Índice local de avatares (fallback quando banco estiver indisponível)
def _avatar_index_path():
    return os.path.join(app.config['UPLOAD_FOLDER'], 'avatars', '_index.json')

def _update_avatar_index(uid: int, rel_path: str) -> None:
    try:
        avatars_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'avatars')
        os.makedirs(avatars_dir, exist_ok=True)
        idx_path = _avatar_index_path()
        data = {}
        if os.path.exists(idx_path):
            try:
                with open(idx_path, 'r', encoding='utf-8') as f:
                    data = json.load(f) or {}
            except Exception:
                data = {}
        data[str(uid)] = {'path': rel_path, 'ts': int(time.time())}
        with open(idx_path, 'w', encoding='utf-8') as f:
            json.dump(data, f)
    except Exception:
        pass

def _read_avatar_index(uid: int) -> str:
    try:
        idx_path = _avatar_index_path()
        if os.path.exists(idx_path):
            with open(idx_path, 'r', encoding='utf-8') as f:
                data = json.load(f) or {}
            obj = data.get(str(uid))
            if obj and isinstance(obj, dict):
                p = obj.get('path') or ''
                return str(p)
    except Exception:
        pass
    return ''

def check_rate_limit(key: str, limit: int = 20, window: int = 60) -> bool:
    now = time.time()
    bucket = RATE_LIMIT.get(key, {'count': 0, 'reset': now + window})
    if now > bucket['reset']:
        bucket = {'count': 0, 'reset': now + window}
    bucket['count'] += 1
    RATE_LIMIT[key] = bucket
    return bucket['count'] <= limit

def audit_log(event: str, details: dict):
    try:
        log_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs', 'audit.log')
        with open(log_path, 'a', encoding='utf-8') as f:
            ts = datetime.now().isoformat()
            user = session.get('user_id')
            ip = request.remote_addr
            f.write(f"{ts}\t{ip}\tuser={user}\t{event}\t{details}\n")
    except Exception:
        pass

def send_reset_email(to_email: str, reset_url: str, token: str | None = None) -> bool:
    host = os.getenv('SMTP_HOST')
    # Porta padrão depende da segurança
    default_port = 587
    security = (os.getenv('SMTP_SECURITY', '').strip().lower() or 'starttls')
    if security == 'ssl':
        default_port = 465
    port = int(os.getenv('SMTP_PORT', str(default_port)))
    user = os.getenv('SMTP_USER')
    password = os.getenv('SMTP_PASSWORD')
    sender = (os.getenv('SMTP_FROM') or user or '').strip()
    sender_name = (os.getenv('SMTP_FROM_NAME') or 'INPROLIB').strip()
    debug_level = int(os.getenv('SMTP_DEBUG', '0'))

    if not host or not user or not password or not sender:
        print('[SMTP] Configuração incompleta. Não foi possível enviar e-mail.')
        print('[SMTP] Dica: configure SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASSWORD e SMTP_FROM.')
        print('[SMTP] Código de redefinição:', token)
        return False

    try:
        msg = EmailMessage()
        msg['Subject'] = 'INPROLIB - Redefinição de senha'
        # Preferir SMTP_FROM (remetente configurado) e incluir nome amigável
        msg['From'] = f"{sender_name} <{sender}>"
        msg['To'] = to_email
        msg.set_content(
            (
                'Olá,\n\nVocê solicitou a redefinição de senha no INPROLIB.\n'
                f'Use o código abaixo para criar uma nova senha (expira em {RESET_TOKEN_EXP_SECONDS} segundos):\n\n{token or "[token indisponível]"}\n\n'
                'Acesse a página "Recuperar senha" e insira o código recebido.\n'
                'Se você não solicitou, ignore este e-mail.'
            )
        )
        
        def _send_via_starttls() -> None:
            with smtplib.SMTP(host, port) as smtp:
                smtp.set_debuglevel(debug_level)
                smtp.ehlo()
                smtp.starttls()
                smtp.ehlo()
                smtp.login(user, password)
                smtp.send_message(msg)

        def _send_via_ssl() -> None:
            with smtplib.SMTP_SSL(host, port) as smtp:
                smtp.set_debuglevel(debug_level)
                smtp.ehlo()
                smtp.login(user, password)
                smtp.send_message(msg)

        def _send_plain() -> None:
            with smtplib.SMTP(host, port) as smtp:
                smtp.set_debuglevel(debug_level)
                smtp.ehlo()
                # Sem STARTTLS
                smtp.login(user, password)
                smtp.send_message(msg)

        # Tenta conforme configuração; depois usa fallbacks inteligentes
        if security == 'ssl':
            try:
                _send_via_ssl()
            except Exception as e1:
                print('[SMTP] SSL falhou, tentando STARTTLS:', e1)
                try:
                    # Ajusta porta padrão quando alterna para STARTTLS
                    if port == 465:
                        _port = 587
                    else:
                        _port = port
                    with smtplib.SMTP(host, _port) as smtp:
                        smtp.set_debuglevel(debug_level)
                        smtp.ehlo(); smtp.starttls(); smtp.ehlo(); smtp.login(user, password); smtp.send_message(msg)
                except Exception as e2:
                    print('[SMTP] STARTTLS falhou, tentando sem TLS:', e2)
                    _send_plain()
        elif security == 'none':
            try:
                _send_plain()
            except Exception as e1:
                print('[SMTP] Login sem TLS falhou, tentando STARTTLS:', e1)
                try:
                    # Se porta comum 25 foi usada, tenta 587
                    _port = 587 if port in (25,) else port
                    with smtplib.SMTP(host, _port) as smtp:
                        smtp.set_debuglevel(debug_level)
                        smtp.ehlo(); smtp.starttls(); smtp.ehlo(); smtp.login(user, password); smtp.send_message(msg)
                except Exception as e2:
                    print('[SMTP] STARTTLS falhou, tentando SSL 465:', e2)
                    with smtplib.SMTP_SSL(host, 465) as smtp:
                        smtp.set_debuglevel(debug_level)
                        smtp.ehlo(); smtp.login(user, password); smtp.send_message(msg)
        else:  # starttls
            try:
                _send_via_starttls()
            except Exception as e1:
                print('[SMTP] STARTTLS falhou, tentando SSL 465:', e1)
                try:
                    with smtplib.SMTP_SSL(host, 465) as smtp:
                        smtp.set_debuglevel(debug_level)
                        smtp.ehlo(); smtp.login(user, password); smtp.send_message(msg)
                except Exception as e2:
                    print('[SMTP] SSL falhou, tentando sem TLS:', e2)
                    _send_plain()

        return True
    except Exception as e:
        print('[SMTP] Erro ao enviar e-mail:', e)
        if 'Username and Password not accepted' in str(e) or '5.7.8' in str(e):
            print('[SMTP] Dica: no Gmail, habilite 2FA e use uma "Senha de app".')
            print('[SMTP] Ajuda: https://support.google.com/accounts/answer/185833')
        if '530 5.7.0' in str(e) or 'STARTTLS' in str(e):
            print('[SMTP] Dica: o servidor requer STARTTLS. Defina SMTP_SECURITY=starttls e SMTP_PORT=587.')
        if 'Must issue a STARTTLS command first' in str(e):
            print('[SMTP] Dica: habilite STARTTLS ou use porta 587.')
        if 'SSL' in str(e) and 'wrong version number' in str(e).lower():
            print('[SMTP] Dica: ajuste SMTP_SECURITY=ssl e SMTP_PORT=465 para servidores que exigem SSL.')
        print('[SMTP] Código de redefinição:', token)
        return False


def send_support_email(body_text: str, attachment: tuple | None = None, reply_to: str | None = None, subject: str = 'INPROLIB - Suporte') -> bool:
    """
    Envia um e-mail de suporte.
    - body_text: texto principal da mensagem
    - attachment: tuple opcional (filename, data_bytes, mimetype)
    - reply_to: e-mail do usuário para facilitar resposta
    """
    host = os.getenv('SMTP_HOST')
    port = int(os.getenv('SMTP_PORT', '587'))
    user = os.getenv('SMTP_USER')
    password = os.getenv('SMTP_PASSWORD')
    sender = os.getenv('SMTP_FROM', user or '')
    use_ssl = os.getenv('SMTP_USE_SSL', '0').lower() in {'1','true','yes'}
    to_email = os.getenv('SUPPORT_EMAIL', 'suporteinprolib@gmail.com')

    if not host or not user or not password or not sender:
        print('[SMTP] Configuração incompleta. Não foi possível enviar e-mail de suporte.')
        return False

    try:
        msg = EmailMessage()
        msg['Subject'] = subject
        msg['From'] = sender
        msg['To'] = to_email
        if reply_to:
            msg['Reply-To'] = reply_to
        msg.set_content(body_text)

        if attachment:
            filename, data_bytes, mimetype = attachment
            try:
                maintype, subtype = (mimetype or 'application/octet-stream').split('/', 1)
            except Exception:
                maintype, subtype = 'application', 'octet-stream'
            msg.add_attachment(data_bytes, maintype=maintype, subtype=subtype, filename=filename)

        if use_ssl:
            with smtplib.SMTP_SSL(host, port) as smtp:
                smtp.login(user, password)
                smtp.send_message(msg)
        else:
            with smtplib.SMTP(host, port) as smtp:
                smtp.ehlo()
                smtp.starttls()
                smtp.ehlo()
                smtp.login(user, password)
                smtp.send_message(msg)
        return True
    except Exception as e:
        print('[SMTP] Erro ao enviar e-mail de suporte:', e)
        return False

def validar_cpf(cpf: str) -> bool:
    if not cpf:
        return False
    cpf = re.sub(r'[^0-9]', '', cpf)
    if len(cpf) != 11:
        return False
    if cpf == cpf[0] * 11:
        return False
    # primeiro dígito verificador
    soma = sum(int(cpf[i]) * (10 - i) for i in range(9))
    resto = 11 - (soma % 11)
    d1 = 0 if resto >= 10 else resto
    # segundo dígito verificador
    soma = sum(int(cpf[i]) * (11 - i) for i in range(10))
    resto = 11 - (soma % 11)
    d2 = 0 if resto >= 10 else resto
    return cpf[-2:] == f"{d1}{d2}"

# Rotas para servir assets conforme os caminhos usados nos HTML (sem alterar HTML/CSS)
@app.route('/<asset_name>.css')
def serve_css(asset_name):
    resp = make_response(send_from_directory(os.path.join(app.static_folder, 'css'), f'{asset_name}.css'))
    # Evita cache agressivo durante desenvolvimento, garantindo atualização imediata do menu
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp

@app.route('/<script_name>.js')
def serve_js(script_name):
    resp = make_response(send_from_directory(os.path.join(app.static_folder, 'javascript'), f'{script_name}.js'))
    # Evita cache agressivo durante desenvolvimento, garantindo atualização imediata do JS
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp

@app.route('/img/<path:filename>')
def serve_img(filename):
    resp = make_response(send_from_directory(os.path.join(app.static_folder, 'img'), filename))
    resp.headers['Cache-Control'] = 'public, max-age=604800'
    return resp

# Política de Senha: entre 8 e 16 caracteres com maiúscula, minúscula, número e símbolo
def password_policy_ok(pwd: str) -> bool:
    try:
        if not isinstance(pwd, str):
            return False
        if len(pwd) < 8 or len(pwd) > 16:
            return False
        return (
            re.search(r'[A-Z]', pwd) and
            re.search(r'[a-z]', pwd) and
            re.search(r'\d', pwd) and
            re.search(r'[^A-Za-z0-9]', pwd)
        ) is not None
    except Exception:
        return False

# Função para conectar ao banco de dados
# Usa psycopg (v3) e cria o banco automaticamente se ele não existir
SCHEMA_READY = False

def get_db_connection():
    """Cria uma conexão com o Postgres usando DATABASE_URL ou DB_CONFIG.
    Em caso de erro, retorna None sem interromper o fluxo da aplicação.
    """
    try:
        db_url = os.getenv('DATABASE_URL')
        # Permite configurar schema via variável de ambiente (padrão: public)
        db_schema = os.getenv('DB_SCHEMA', 'public').strip() or 'public'
        if db_url:
            if db_url.startswith("postgres://"):
                db_url = db_url.replace("postgres://", "postgresql://", 1)
            # Timeout curto para evitar travamentos quando o serviço está indisponível
            conn = psycopg.connect(db_url, connect_timeout=5)
            # Garante search_path correto para o schema informado
            try:
                with conn.cursor() as cur:
                    cur.execute(f'SET search_path TO "{db_schema}", public')
            except Exception:
                pass
            return conn
        # Sanitiza valores do DB_CONFIG (remove espaços e normaliza porta)
        cfg = {k: (str(v).strip() if isinstance(v, str) else v) for k, v in DB_CONFIG.items()}
        try:
            # Porta deve ser int para algumas instalações
            if 'port' in cfg:
                try:
                    cfg['port'] = int(str(cfg['port']).strip())
                except Exception:
                    # Fallback: mantém como string se conversão falhar
                    pass
        except Exception:
            pass

        conn = psycopg.connect(**cfg, connect_timeout=5)
        try:
            with conn.cursor() as cur:
                cur.execute(f'SET search_path TO "{db_schema}", public')
        except Exception:
            pass
        return conn
    except InvalidCatalogName:
        # Tenta criar o banco se não existir e reconecta
        try:
            tmp = {k: (str(v).strip() if isinstance(v, str) else v) for k, v in DB_CONFIG.items()}
            dbname = tmp.pop('dbname', None)
            # Garante tipo correto da porta
            if 'port' in tmp:
                try:
                    tmp['port'] = int(str(tmp['port']).strip())
                except Exception:
                    pass
            # Conecta no banco administrativo padrão 'postgres'
            admin = psycopg.connect(**{**tmp, 'dbname': 'postgres'}, connect_timeout=5)
            cur = admin.cursor()
            if dbname:
                cur.execute(f'CREATE DATABASE "{dbname}"')
                admin.commit()
            cur.close(); admin.close()
            conn = psycopg.connect(**DB_CONFIG, connect_timeout=5)
            try:
                with conn.cursor() as cur2:
                    _schema = os.getenv("DB_SCHEMA", "public").strip() or "public"
                    cur2.execute(f'SET search_path TO "{_schema}", public')
            except Exception:
                pass
            return conn
        except Exception as e:
            # Log detalhado para facilitar diagnóstico
            print(f"[DB] Falha ao criar/conectar banco: {e}")
            return None
    except Exception as e:
        # Se o erro indicar banco inexistente, tenta criar automaticamente
        msg = str(e).lower()
        # Cobrir casos de encoding quebrado: 'não existe' pode aparecer como 'n�o existe' em alguns logs
        if (
            'does not exist' in msg or
            'não existe' in msg or  # UTF-8 correto
            'n�o existe' in msg or  # variante com REPLACEMENT CHARACTER
            'invalidcatalogname' in msg
        ):
            try:
                tmp = {k: (str(v).strip() if isinstance(v, str) else v) for k, v in DB_CONFIG.items()}
                dbname = tmp.pop('dbname', None)
                if 'port' in tmp:
                    try:
                        tmp['port'] = int(str(tmp['port']).strip())
                    except Exception:
                        pass
                admin = psycopg.connect(**{**tmp, 'dbname': 'postgres'}, connect_timeout=5)
                with admin.cursor() as cur:
                    if dbname:
                        cur.execute(f'CREATE DATABASE "{dbname}"')
                        admin.commit()
                admin.close()
                conn = psycopg.connect(**DB_CONFIG, connect_timeout=5)
                try:
                    with conn.cursor() as cur2:
                        _schema = os.getenv("DB_SCHEMA", "public").strip() or "public"
                        cur2.execute(f'SET search_path TO "{_schema}", public')
                    
                except Exception:
                    pass
                return conn
            except Exception as e2:
                print(f"[DB] Tentativa de criação do banco falhou: {e2}")
                return None
        # Loga erro real (ex.: connection refused, auth failed, etc.)
        print(f"[DB] Erro ao conectar ao banco de dados: {e}")
        return None

# Cache simples para rótulos de enum de status de publicação
STATUS_LABEL_CACHE: dict[str, str] = {}

def _norm(s: str) -> str:
    try:
        return ''.join(c for c in unicodedata.normalize('NFKD', s or '') if not unicodedata.combining(c)).lower().strip()
    except Exception:
        return (s or '').lower().strip()

def status_label(preferred: str) -> str:
    """Resolve o rótulo correto no enum 'status_publicacao', tolerando variações.
    Retorna o rótulo existente equivalente ao preferred (case-insensitive e sem acentos),
    ou o primeiro rótulo disponível como fallback.
    """
    key = _norm(preferred)
    if key in STATUS_LABEL_CACHE:
        return STATUS_LABEL_CACHE[key]
    conn = get_db_connection()
    if not conn:
        return preferred
    try:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT e.enumlabel
            FROM pg_enum e
            JOIN pg_type t ON e.enumtypid = t.oid
            WHERE t.typname = %s
            ORDER BY e.enumsortorder
            """,
            ('status_publicacao',)
        )
        labels = [r[0] for r in cur.fetchall()]
        cur.close(); conn.close()
        # match exato por normalização
        for lab in labels:
            if _norm(lab) == key:
                STATUS_LABEL_CACHE[key] = lab
                return lab
        # heurísticas por semântica
        if key in {'publicado','publicada'}:
            for lab in labels:
                if 'public' in _norm(lab):
                    STATUS_LABEL_CACHE[key] = lab
                    return lab
        if key in {'reprovado','reprovada','indeferido','indeferida'}:
            for lab in labels:
                nl = _norm(lab)
                if 'reprov' in nl or 'indef' in nl:
                    STATUS_LABEL_CACHE[key] = lab
                    return lab
            # Se não encontrou rótulos equivalentes, tenta garantir a existência de um rótulo de reprovação
            try:
                ensure_status_enum_reprovado()
            except Exception:
                pass
            # Reconsulta labels após tentativa de garantir o rótulo
            try:
                conn2 = get_db_connection()
                if conn2:
                    cur2 = conn2.cursor()
                    cur2.execute(
                        """
                        SELECT e.enumlabel
                        FROM pg_enum e
                        JOIN pg_type t ON e.enumtypid = t.oid
                        WHERE t.typname = %s
                        ORDER BY e.enumsortorder
                        """,
                        ('status_publicacao',)
                    )
                    labels2 = [r[0] for r in cur2.fetchall()]
                    cur2.close(); conn2.close()
                    for lab in labels2:
                        nl = _norm(lab)
                        if 'reprov' in nl or 'indef' in nl:
                            STATUS_LABEL_CACHE[key] = lab
                            return lab
            except Exception:
                pass
            # Fallback seguro para pedidos de reprovação: evita cair em 'Publicado'
            for lab in labels:
                nl = _norm(lab)
                if 'pend' in nl or 'avali' in nl or 'analis' in nl:
                    STATUS_LABEL_CACHE[key] = lab
                    return lab
        # suporte explícito ao rótulo "Denunciado"
        if key in {'denunciado','denunciada'}:
            for lab in labels:
                nl = _norm(lab)
                if 'denunc' in nl:
                    STATUS_LABEL_CACHE[key] = lab
                    return lab
        if key == 'pendente':
            # 1) Tentativa direta por "pend"
            for lab in labels:
                if 'pend' in _norm(lab):
                    STATUS_LABEL_CACHE[key] = lab
                    return lab
            # 2) Heurística para rótulos de análise/avaliação
            for lab in labels:
                nl = _norm(lab)
                if ('avali' in nl or 'analis' in nl or 'em ' in nl) and 'public' not in nl and 'reprov' not in nl and 'indef' not in nl:
                    STATUS_LABEL_CACHE[key] = lab
                    return lab
            # 3) Fallback seguro: escolher rótulo que não indique publicação nem reprovação
            for lab in labels:
                nl = _norm(lab)
                if 'public' not in nl and 'reprov' not in nl and 'indef' not in nl:
                    STATUS_LABEL_CACHE[key] = lab
                    return lab
        # fallback: primeiro disponível
        if labels:
            STATUS_LABEL_CACHE[key] = labels[0]
            return labels[0]
        return preferred
    except Exception:
        try:
            conn.close()
        except Exception:
            pass
        return preferred

        # Monta a conexão via dict (DB_CONFIG)
        cfg = {k: (str(v).strip() if isinstance(v, str) else v) for k, v in DB_CONFIG.items()}
        try:
            conn = psycopg.connect(**cfg)
            return conn
        except InvalidCatalogName:
            # Banco não existe: conecta no 'postgres' e cria
            admin_cfg = {**cfg, 'dbname': 'postgres'}
            admin = psycopg.connect(**admin_cfg)
            admin.autocommit = True
            with admin.cursor() as cur:
                cur.execute("SELECT 1 FROM pg_database WHERE datname=%s", (cfg['dbname'],))
                if not cur.fetchone():
                    cur.execute(f'CREATE DATABASE "{cfg["dbname"]}"')
            admin.close()
            # Conecta ao banco recém-criado
            conn = psycopg.connect(**cfg)
            return conn
    except Exception as e:
        print(f"Erro ao conectar ao banco de dados: {e}")
        return None

# Helper para garantir coluna 'ativo' em curso
def ensure_curso_ativo_column():
    conn = get_db_connection()
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute("SELECT 1 FROM information_schema.columns WHERE table_name='curso' AND column_name='ativo'")
        if not cur.fetchone():
            cur.execute("ALTER TABLE curso ADD COLUMN ativo BOOLEAN NOT NULL DEFAULT TRUE")
            conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        print(f"Falha ao garantir coluna curso.ativo: {e}")

# Helper para garantir coluna 'ativo' em usuario
def ensure_usuario_ativo_column():
    conn = get_db_connection()
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute("SELECT current_schema()")
        schema = (cur.fetchone() or ['public'])[0] or 'public'
        cur.execute("SELECT 1 FROM information_schema.columns WHERE table_schema=%s AND table_name='usuario' AND column_name='ativo'", (schema,))
        if not cur.fetchone():
            cur.execute(f'ALTER TABLE "{schema}".usuario ADD COLUMN ativo BOOLEAN NOT NULL DEFAULT TRUE')
            conn.commit()
        cur.close(); conn.close()
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        print(f"Falha ao garantir coluna usuario.ativo: {e}")

# Helper para garantir colunas de endereço em usuario (cep, logradouro, complemento, bairro, cidade, estado)
def ensure_usuario_endereco_columns():
    conn = get_db_connection()
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute("SELECT current_schema()")
        schema = (cur.fetchone() or ['public'])[0] or 'public'
        cols = [
            ('cep', 'VARCHAR(9)'),
            ('logradouro', 'VARCHAR(255)'),
            ('complemento', 'VARCHAR(255)'),
            ('bairro', 'VARCHAR(255)'),
            ('cidade', 'VARCHAR(255)'),
            ('estado', 'VARCHAR(2)')
        ]
        for col, coldef in cols:
            cur.execute("SELECT 1 FROM information_schema.columns WHERE table_schema=%s AND table_name='usuario' AND column_name=%s", (schema, col))
            if not cur.fetchone():
                cur.execute(f'ALTER TABLE "{schema}".usuario ADD COLUMN {col} {coldef}')
                conn.commit()
        cur.close(); conn.close()
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        print(f"Falha ao garantir colunas de endereço em usuario: {e}")

# Helper para garantir colunas de preferências do usuário (telefone, tema_preferido)
def ensure_usuario_preferences_columns():
    conn = get_db_connection()
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute("SELECT current_schema()")
        schema = (cur.fetchone() or ['public'])[0] or 'public'
        prefs = [
            ('telefone', 'VARCHAR(20)'),
            ('tema_preferido', "VARCHAR(16) DEFAULT 'claro'")
        ]
        for col, coldef in prefs:
            cur.execute("SELECT 1 FROM information_schema.columns WHERE table_schema=%s AND table_name='usuario' AND column_name=%s", (schema, col))
            if not cur.fetchone():
                cur.execute(f'ALTER TABLE "{schema}".usuario ADD COLUMN {col} {coldef}')
                conn.commit()
        cur.close(); conn.close()
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        print(f"Falha ao garantir colunas de preferências em usuario: {e}")

# Helper para garantir existência da tabela 'avaliacao'
def ensure_avaliacao_table():
    conn = get_db_connection()
    if not conn:
        return
    try:
        cur = conn.cursor()
        # Verifica a existência da tabela no search_path atual
        cur.execute("SELECT to_regclass('avaliacao')")
        reg = cur.fetchone()
        exists = reg and reg[0] is not None
        if not exists:
            # Cria com estrutura compatível ao banco.sql
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS avaliacao (
                    id_avaliacao SERIAL PRIMARY KEY,
                    id_publicacao INTEGER NOT NULL REFERENCES publicacao(id_publicacao),
                    id_avaliador INTEGER NOT NULL REFERENCES usuario(id_usuario),
                    nota NUMERIC(3,1),
                    comentario TEXT,
                    data_avaliacao TIMESTAMP WITHOUT TIME ZONE DEFAULT CURRENT_TIMESTAMP
                )
                """
            )
            conn.commit()
        cur.close(); conn.close()
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        print(f"Falha ao garantir tabela avaliacao: {e}")

# Garante que o enum de status possua um rótulo de pendência (e.g., Pendente/Em avaliação)
def ensure_status_enum_pending():
    conn = get_db_connection()
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT e.enumlabel
            FROM pg_enum e
            JOIN pg_type t ON e.enumtypid = t.oid
            WHERE t.typname = %s
            ORDER BY e.enumsortorder
            """,
            ('status_publicacao',)
        )
        labels = [r[0] for r in cur.fetchall()]
        has_pending = False
        for lab in labels:
            nl = _norm(lab)
            if 'pend' in nl or 'avali' in nl or 'analis' in nl:
                has_pending = True
                break
        if not has_pending:
            try:
                cur.execute("ALTER TYPE status_publicacao ADD VALUE 'Pendente'")
                conn.commit()
            except Exception:
                # Ignora se já existir ou não puder adicionar
                pass
        cur.close(); conn.close()
    except Exception:
        try:
            conn.close()
        except Exception:
            pass

# Garante que o enum de status possua um rótulo para denúncias (e.g., Denunciado)
def ensure_status_enum_denunciado():
    conn = get_db_connection()
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT e.enumlabel
            FROM pg_enum e
            JOIN pg_type t ON e.enumtypid = t.oid
            WHERE t.typname = %s
            ORDER BY e.enumsortorder
            """,
            ('status_publicacao',)
        )
        labels = [r[0] for r in cur.fetchall()]
        has_denunc = any('denunc' in _norm(lab) for lab in labels)
        if not has_denunc:
            try:
                cur.execute("ALTER TYPE status_publicacao ADD VALUE 'Denunciado'")
                conn.commit()
            except Exception:
                # Ignora se já existir ou não puder adicionar
                pass
        cur.close(); conn.close()
    except Exception:
        try:
            conn.close()
        except Exception:
            pass

# Garante que o enum de status possua um rótulo de reprovação (e.g., Reprovado/Indeferido)
def ensure_status_enum_reprovado():
    conn = get_db_connection()
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT e.enumlabel
            FROM pg_enum e
            JOIN pg_type t ON e.enumtypid = t.oid
            WHERE t.typname = %s
            ORDER BY e.enumsortorder
            """,
            ('status_publicacao',)
        )
        labels = [r[0] for r in cur.fetchall()]
        has_reprov = any(('reprov' in _norm(lab)) or ('indef' in _norm(lab)) for lab in labels)
        if not has_reprov:
            try:
                cur.execute("ALTER TYPE status_publicacao ADD VALUE 'Reprovado'")
                conn.commit()
            except Exception:
                # Ignora se já existir ou não puder adicionar
                pass
        cur.close(); conn.close()
    except Exception:
        try:
            conn.close()
        except Exception:
            pass

# Helper para garantir coluna de vínculo direto usuario.id_curso_usuario
def ensure_usuario_curso_fk_column():
    conn = get_db_connection()
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute("SELECT current_schema()")
        schema = (cur.fetchone() or ['public'])[0] or 'public'
        # verifica se existe coluna id_curso_usuario
        cur.execute(
            "SELECT 1 FROM information_schema.columns WHERE table_schema=%s AND table_name='usuario' AND column_name='id_curso_usuario'",
            (schema,)
        )
        if not cur.fetchone():
            cur.execute(f'ALTER TABLE "{schema}".usuario ADD COLUMN id_curso_usuario INTEGER NULL')
            conn.commit()
            # tenta criar FK
            try:
                cur.execute(
                    f'ALTER TABLE "{schema}".usuario ADD CONSTRAINT fk_usuario_curso FOREIGN KEY (id_curso_usuario) REFERENCES "{schema}".curso(id_curso) ON DELETE SET NULL'
                )
                conn.commit()
            except Exception:
                pass
        cur.close(); conn.close()
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        print(f"Falha ao garantir coluna usuario.id_curso_usuario: {e}")

# Helper para garantir coluna 'id_orientador' em publicacao
def ensure_publicacao_orientador_column():
    conn = get_db_connection()
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute("SELECT current_schema()")
        schema = (cur.fetchone() or ['public'])[0] or 'public'
        # adiciona coluna se não existir
        cur.execute("SELECT 1 FROM information_schema.columns WHERE table_schema=%s AND table_name='publicacao' AND column_name='id_orientador'", (schema,))
        if not cur.fetchone():
            cur.execute(f'ALTER TABLE "{schema}".publicacao ADD COLUMN id_orientador INTEGER NULL')
            conn.commit()
        # tenta criar FK (ignora erro se já existir)
        try:
            cur.execute(f'ALTER TABLE "{schema}".publicacao ADD CONSTRAINT fk_publicacao_orientador FOREIGN KEY (id_orientador) REFERENCES "{schema}".usuario(id_usuario) ON DELETE SET NULL')
            conn.commit()
        except Exception:
            pass
        cur.close(); conn.close()
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        print(f"Falha ao garantir coluna publicacao.id_orientador: {e}")

# Inicialização do schema na primeira requisição
SCHEMA_INIT_DONE = False
@app.before_request
def init_schema_once():
    global SCHEMA_INIT_DONE
    if SCHEMA_INIT_DONE:
        return
    try:
        ensure_usuario_ativo_column()
        ensure_usuario_endereco_columns()
        ensure_publicacao_orientador_column()
        ensure_usuario_curso_fk_column()
        ensure_usuario_preferences_columns()
        ensure_avaliacao_table()
        # garante rótulos essenciais no enum de status
        try:
            ensure_status_enum_pending()
        except Exception:
            pass
        try:
            ensure_status_enum_denunciado()
        except Exception:
            pass
        try:
            ensure_status_enum_reprovado()
        except Exception:
            pass
        SCHEMA_INIT_DONE = True
        print("Schema inicial garantido: usuario.ativo, endereço, publicacao.id_orientador, usuario.id_curso_usuario, preferências e tabela avaliacao.")
    except Exception as e:
        print(f"Falha ao garantir schema inicial: {e}")

# Decorator para verificar se o usuário está logado
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Por favor, faça login para acessar esta página', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# Controle de acesso por perfil
def _normalize_role(tipo: str) -> str:
    if not tipo:
        return ''
    tipo = str(tipo).strip()
    mapping = {
        'Administrador': 'Administrador',
        'Funcionário': 'Administrador',
        'Funcionario': 'Administrador',
        'Professor': 'Docente',
        'Docente': 'Docente',
        'Aluno': 'Aluno'
    }
    return mapping.get(tipo, tipo)

def roles_required(allowed_roles):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            role = _normalize_role(session.get('role') or session.get('user_tipo'))
            if not role:
                flash('Por favor, faça login para acessar esta página', 'error')
                return redirect(url_for('login'))
            if role not in allowed_roles:
                # Silenciosamente redireciona para Home sem exibir mensagem
                return redirect(url_for('home'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator

# Rota principal -> redireciona para Home
@app.route('/')
def index():
    return redirect(url_for('login'))

# Rota segura para cadastrar/atualizar usuário administrador
@app.route('/setup_admin', methods=['GET'])
def setup_admin():
    try:
        token = (request.args.get('token') or '').strip()
        if token != ADMIN_SETUP_TOKEN:
            return make_response(jsonify({'error': 'Unauthorized'}), 403)

        # Dados padrão do admin (podem ser sobrepostos via querystring)
        nome = (request.args.get('nome') or 'Samuel Edgar').strip()
        email = (request.args.get('email') or 'samuel.edgar@gmail.com').strip()
        cpf = (request.args.get('cpf') or '000.000.000-00').strip()
        # Permitir sobrepor a senha via querystring ("senha" ou "password"); fallback para ADMIN_TEMP_PASSWORD
        senha_param = (request.args.get('senha') or request.args.get('password') or '').strip()
        temp_password = senha_param if senha_param else ADMIN_TEMP_PASSWORD
        senha_hash = generate_password_hash(temp_password)

        conn = get_db_connection()
        if not conn:
            return make_response(jsonify({'error': 'Falha ao conectar ao banco.'}), 500)
        try:
            cur = conn.cursor(row_factory=dict_row)
            cur.execute("SELECT id_usuario, tipo FROM usuario WHERE email = %s", (email,))
            user = cur.fetchone()

            # Detectar o label válido para 'Funcionário' no enum (com ou sem acento)
            cur_labels = conn.cursor()
            cur_labels.execute("""
                SELECT e.enumlabel
                FROM pg_enum e
                WHERE e.enumtypid = (
                  SELECT a.atttypid
                  FROM pg_attribute a
                  JOIN pg_class c ON a.attrelid = c.oid
                  WHERE c.relname = 'usuario' AND a.attname = 'tipo'
                  LIMIT 1
                )
            """)
            rows = cur_labels.fetchall() or []
            labels = set()
            for r in rows:
                try:
                    labels.add(r[0])
                except Exception:
                    try:
                        labels.add(r.get('enumlabel'))
                    except Exception:
                        pass
            cur_labels.close()
            tipo_admin_label = 'Funcionário' if 'Funcionário' in labels else ('Funcionario' if 'Funcionario' in labels else None)
            if not tipo_admin_label:
                audit_log('setup_admin_error', {'error': 'Enum tipo_usuario sem Funcionário/Funcionario', 'labels': sorted(list(labels))})
                cur.close(); conn.close()
                return make_response(jsonify({'error': f"Enum tipo_usuario não possui 'Funcionário' ou 'Funcionario'. Labels: {sorted(list(labels))}"}), 500)

            if user:
                cur2 = conn.cursor()
                cur2.execute(
                    "UPDATE usuario SET nome = %s, cpf = %s, senha = %s, tipo = %s WHERE id_usuario = %s",
                    (nome, cpf, senha_hash, tipo_admin_label, user['id_usuario'])
                )
                conn.commit()
                cur2.close()
                status = 'updated'
            else:
                cur2 = conn.cursor()
                cur2.execute(
                    "INSERT INTO usuario (nome, email, cpf, senha, tipo, curso_usuario) VALUES (%s, %s, %s, %s, %s, %s)",
                    (nome, email, cpf, senha_hash, tipo_admin_label, None)
                )
                conn.commit()
                cur2.close()
                status = 'created'

            cur.close()
            conn.close()
            audit_log('setup_admin_ok', {'email': email, 'status': status})
            return jsonify({'ok': True, 'status': status, 'email': email, 'temp_password': temp_password})
        except Exception as e:
            try:
                conn.close()
            except Exception:
                pass
            audit_log('setup_admin_error', {'error': str(e)})
            return make_response(jsonify({'error': str(e)}), 500)
    except Exception as e:
        audit_log('setup_admin_error_outer', {'error': str(e)})
        return make_response(jsonify({'error': str(e)}), 500)

# Rota para a página inicial após login
@app.route('/home')
@login_required
def home():
    conn = get_db_connection()
    publicacoes = []
    
    if conn:
        try:
            cur = conn.cursor(row_factory=dict_row)
            # Buscar as últimas publicações
            cur.execute("""
                SELECT p.*, u.nome as autor_nome, c.nome_curso 
                FROM publicacao p
                JOIN usuario u ON p.id_autor = u.id_usuario
                JOIN curso c ON p.id_curso = c.id_curso
                WHERE p.status = %s
                ORDER BY p.data_publicacao DESC
                LIMIT 10
            """, (status_label('Publicado'),))
            publicacoes = cur.fetchall()
            cur.close()
            conn.close()
        except Exception as e:
            flash(f'Erro ao buscar publicações: {e}', 'error')
    
    return render_template('home.html', publicacoes=publicacoes)

# (Removida) Rota de cadastro de login

# Tela de Login
@app.route('/login', methods=['GET', 'POST'])
def login():
    # Auto-login: se já houver sessão ativa, redireciona para home
    if request.method == 'GET' and session.get('user_id'):
        return redirect(url_for('home'))
    if request.method == 'POST':
        key = f"{request.remote_addr}:login"
        if not check_rate_limit(key, limit=30, window=60):
            flash('Muitas tentativas. Tente novamente em instantes.', 'error')
            audit_log('rate_limit', {'route': 'login'})
            return redirect(url_for('login'))

        email = (request.form.get('email') or '').strip()
        cpf = (request.form.get('cpf') or '').strip()
        senha = (request.form.get('senha') or '').strip()

        if not senha or (not email and not cpf):
            flash('Informe seu CPF (ou e-mail) e a senha.', 'error')
            return redirect(url_for('login'))

        if cpf and not validar_cpf(cpf):
            flash('CPF inválido.', 'error')
            return redirect(url_for('login'))

        conn = get_db_connection()
        if not conn:
            flash('Falha ao conectar ao banco.', 'error')
            return redirect(url_for('login'))
        try:
            cur = conn.cursor(row_factory=dict_row)
            if email:
                cur.execute("SELECT id_usuario, nome, email, senha, tipo, foto_perfil FROM usuario WHERE email = %s", (email,))
            else:
                cpf_digits = re.sub(r'[^0-9]', '', cpf)
                cur.execute("SELECT id_usuario, nome, email, senha, tipo, foto_perfil FROM usuario WHERE regexp_replace(cpf, '[^0-9]', '', 'g') = %s", (cpf_digits,))
            user = cur.fetchone()
            cur.close()
            conn.close()
            if not user:
                flash('Usuário não encontrado.', 'error')
                return redirect(url_for('login'))
            senha_hash = user['senha'] if isinstance(user['senha'], str) else (user['senha'].decode() if user['senha'] else '')
            if not senha_hash or not check_password_hash(senha_hash, senha):
                flash('Senha incorreta.', 'error')
                return redirect(url_for('login'))
            session['user_id'] = user['id_usuario']
            session['user_name'] = user['nome']
            session['user_tipo'] = user['tipo']
            session['role'] = _normalize_role(user['tipo'])
            # Lembrar-me: tornar sessão permanente se marcado
            remember_flag = (request.form.get('remember') or '').strip()
            session.permanent = bool(remember_flag)

            # Avatar e tema preferido
            foto = user.get('foto_perfil')
            def _norm_photo_path(fp: str) -> str:
                if not fp:
                    return ''
                fp = str(fp).replace('\\', '/').strip()
                idx = fp.lower().find('static/')
                if idx != -1:
                    rel = fp[idx+len('static/'):]
                    return rel
                if fp.startswith('uploads/'):
                    return fp
                return ''
            rel_photo = _norm_photo_path(foto)
            if not rel_photo:
                # tenta índice local (fallback quando upload ocorreu sem DB)
                try:
                    fallback_rel = _read_avatar_index(user['id_usuario'])
                except Exception:
                    fallback_rel = ''
                rel_photo = _norm_photo_path(fallback_rel) or fallback_rel
            session['user_photo'] = rel_photo

            # carrega tema preferido do usuário
            try:
                theme = None
                conn2 = get_db_connection()
                if conn2:
                    cur2 = conn2.cursor()
                    cur2.execute("SELECT tema_preferido FROM usuario WHERE id_usuario = %s", (user['id_usuario'],))
                    r2 = cur2.fetchone()
                    if r2:
                        theme = r2[0]
                    cur2.close(); conn2.close()
                session['user_theme'] = (theme or 'claro')
            except Exception:
                session['user_theme'] = 'claro'

            audit_log('login_ok', {'email': email or '', 'cpf': cpf or ''})
            return redirect(url_for('home'), code=303)
        except Exception as e:
            try:
                conn.close()
            except Exception:
                pass
            flash(f'Erro no login: {e}', 'error')
            audit_log('login_error', {'error': str(e)})
            return redirect(url_for('login'))

    # GET — gerar captcha para o formulário de cadastro no modal
    a, b = random.randint(1, 9), random.randint(1, 9)
    session['captcha_answer'] = str(a + b)
    captcha_question = f"Quanto é {a} + {b}?"
    return render_template('login.html', captcha_question=captcha_question)

# Upload de avatar do usuário logado
@app.route('/upload_avatar', methods=['POST'])
@login_required
def upload_avatar():
    file = request.files.get('avatar')
    if not file or not file.filename:
        return jsonify({'ok': False, 'error': 'Nenhum arquivo selecionado'}), 400

    # valida extensão
    filename = secure_filename(file.filename)
    name, ext = os.path.splitext(filename)
    ext = ext.lower()
    allowed = {'.png', '.jpg', '.jpeg', '.webp', '.gif'}
    if ext not in allowed:
        return jsonify({'ok': False, 'error': 'Formato não suportado'}), 400

    # diretório de avatares
    avatars_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'avatars')
    os.makedirs(avatars_dir, exist_ok=True)

    # nome do arquivo por usuário
    uid = session.get('user_id')
    ts = int(time.time())
    out_name = f"avatar_{uid}_{ts}{ext}"
    out_path = os.path.join(avatars_dir, out_name)
    try:
        file.save(out_path)
    except Exception as e:
        return jsonify({'ok': False, 'error': f'Falha ao salvar: {e}'}), 500

    # caminho relativo para servir via static
    rel_path = f"uploads/avatars/{out_name}"
    # Atualiza índice local (fallback sem banco)
    try:
        _update_avatar_index(uid, rel_path)
    except Exception:
        pass

    # Atualiza no banco
    db_saved = False
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor()
            cur.execute("UPDATE usuario SET foto_perfil = %s WHERE id_usuario = %s", (rel_path, uid))
            conn.commit()
            cur.close()
            conn.close()
            db_saved = True
        except Exception:
            try:
                conn.close()
            except Exception:
                pass
            # segue com índice local e sessão
    else:
        # sem conexão, segue com índice local e sessão
        pass

    # Atualiza sessão
    session['user_photo'] = rel_path

    # URL acessível
    photo_url = url_for('static', filename=rel_path)
    return jsonify({'ok': True, 'photo_url': photo_url, 'db_saved': db_saved})

# Esqueci a senha: solicitar token
@app.route('/esqueci_senha', methods=['GET', 'POST'])
def esqueci_senha():
    if request.method == 'POST':
        email = (request.form.get('email') or '').strip()
        # Caso o front envie via AJAX, responder JSON em vez de redirect
        if not email:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'ok': False, 'error': 'Informe seu e-mail.'}), 400
            flash('Informe seu e-mail.', 'error')
            return redirect(url_for('esqueci_senha'))

        conn = get_db_connection()
        if not conn:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'ok': False, 'error': 'Falha ao conectar ao banco.'}), 500
            flash('Falha ao conectar ao banco.', 'error')
            return redirect(url_for('esqueci_senha'))
        try:
            cur = conn.cursor()
            # Garantir validação pelo cadastro: busca case-insensitive e usa e-mail canônico do banco
            norm_email = email.lower()
            cur.execute("SELECT email FROM usuario WHERE LOWER(email) = %s", (norm_email,))
            row = cur.fetchone()
            if not row:
                cur.close(); conn.close()
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'ok': False, 'error': 'E-mail não cadastrado.'}), 404
                flash('E-mail não cadastrado.', 'error')
                return redirect(url_for('esqueci_senha'))
            canonical_email = row[0]

            # Expira tokens anteriores vinculados a este e-mail
            cur.execute("UPDATE esqueci_senha SET status = 'Expirado' WHERE email = %s AND status = 'Ativo'", (canonical_email,))
            # Gera novo token numérico (6 dígitos) para fácil digitação
            token = f"{random.randint(100000, 999999)}"
            cur.execute(
                "INSERT INTO esqueci_senha (email, token, data_solicitacao, status) VALUES (%s, %s, %s, %s)",
                (canonical_email, token, datetime.utcnow(), 'Ativo')
            )
            conn.commit()
            cur.close(); conn.close()

            reset_url = url_for('resetar_senha', token=token, email=canonical_email, _external=True)
            send_ok = send_reset_email(canonical_email, reset_url, token)
            audit_log('forgot_ok', {'email': canonical_email, 'email_sent': bool(send_ok)})

            # Suporte a AJAX: se requisitado via XHR, retorna JSON
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                payload = {'ok': True, 'email_sent': bool(send_ok)}
                if not send_ok:
                    payload['error'] = 'Falha ao enviar e-mail. Verifique a configuração SMTP.'
                    # Fallback: incluir token apenas quando explicitamente habilitado
                    show_token = os.getenv('SHOW_RESET_TOKEN_IN_UI', '0').lower() in {'1','true','yes'}
                    if show_token:
                        payload['dev_token'] = token
                return jsonify(payload)
            # Fluxo tradicional: ajusta mensagem conforme resultado
            if send_ok:
                flash('Código enviado! Verifique seu e-mail.', 'success')
            else:
                show_token = os.getenv('SHOW_RESET_TOKEN_IN_UI', '0').lower() in {'1','true','yes'}
                if show_token:
                    flash(f'Falha ao enviar e-mail. Código: {token}', 'info')
                else:
                    flash('Falha ao enviar e-mail. Verifique a configuração SMTP.', 'error')
            return redirect(url_for('esqueci_senha'))
        except Exception as e:
            try:
                conn.close()
            except Exception:
                pass
            audit_log('forgot_error', {'error': str(e)})
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'ok': False, 'error': f'Erro ao gerar token: {e}'}), 500
            flash(f'Erro ao gerar token: {e}', 'error')
            return redirect(url_for('esqueci_senha'))

    return render_template('esqueci_senha.html')

# Endpoint de diagnóstico SMTP
@app.route('/api/smtp/self_test', methods=['GET'])
def smtp_self_test():
    host = os.getenv('SMTP_HOST')
    security = (os.getenv('SMTP_SECURITY', '').strip().lower() or 'starttls')
    default_port = 587 if security != 'ssl' else 465
    port = int(os.getenv('SMTP_PORT', str(default_port)))
    user = os.getenv('SMTP_USER')
    password = os.getenv('SMTP_PASSWORD')
    debug_level = int(os.getenv('SMTP_DEBUG', '0'))

    if not host or not user or not password:
        return jsonify({
            'ok': False,
            'error': 'Configuração incompleta: defina SMTP_HOST, SMTP_PORT, SMTP_USER e SMTP_PASSWORD.'
        }), 400

    try:
        if security == 'ssl':
            try:
                with smtplib.SMTP_SSL(host, port) as smtp:
                    smtp.set_debuglevel(debug_level)
                    smtp.ehlo(); smtp.login(user, password); code, _ = smtp.noop()
            except Exception as e1:
                with smtplib.SMTP(host, 587) as smtp:
                    smtp.set_debuglevel(debug_level)
                    smtp.ehlo(); smtp.starttls(); smtp.ehlo(); smtp.login(user, password); code, _ = smtp.noop()
        elif security == 'none':
            try:
                with smtplib.SMTP(host, port) as smtp:
                    smtp.set_debuglevel(debug_level)
                    smtp.ehlo(); smtp.login(user, password); code, _ = smtp.noop()
            except Exception as e1:
                with smtplib.SMTP(host, 587) as smtp:
                    smtp.set_debuglevel(debug_level)
                    smtp.ehlo(); smtp.starttls(); smtp.ehlo(); smtp.login(user, password); code, _ = smtp.noop()
        else:  # starttls
            try:
                with smtplib.SMTP(host, port) as smtp:
                    smtp.set_debuglevel(debug_level)
                    smtp.ehlo(); smtp.starttls(); smtp.ehlo(); smtp.login(user, password); code, _ = smtp.noop()
            except Exception as e1:
                with smtplib.SMTP_SSL(host, 465) as smtp:
                    smtp.set_debuglevel(debug_level)
                    smtp.ehlo(); smtp.login(user, password); code, _ = smtp.noop()
        return jsonify({'ok': True, 'message': 'Login SMTP OK', 'noop_code': code})
    except Exception as e:
        msg = str(e)
        hint = None
        if 'Username and Password not accepted' in msg or '5.7.8' in msg:
            hint = 'Gmail exige 2FA e "Senha de app". Gere uma senha de app e use em SMTP_PASSWORD.'
        return jsonify({'ok': False, 'error': msg, 'hint': hint}), 500

# Resetar senha via token
@app.route('/resetar_senha', methods=['GET', 'POST'])
def resetar_senha():
    token = request.args.get('token') or request.form.get('token')
    email = request.args.get('email') or request.form.get('email')
    if request.method == 'POST':
        nova = (request.form.get('nova_senha') or '').strip()
        confirmar = (request.form.get('confirmar_senha') or '').strip()
        if not nova or not confirmar or nova != confirmar:
            flash('As senhas devem coincidir e não podem ser vazias.', 'error')
            return redirect(url_for('resetar_senha', token=token, email=email))
        if not password_policy_ok(nova):
            flash('A nova senha deve ter exatamente 8 caracteres, com maiúscula, minúscula, número e símbolo.', 'error')
            return redirect(url_for('resetar_senha', token=token, email=email))
        conn = get_db_connection()
        if not conn:
            flash('Falha ao conectar ao banco.', 'error')
            return redirect(url_for('resetar_senha', token=token, email=email))
        try:
            cur = conn.cursor(row_factory=dict_row)
            # Resolve e-mail canônico e valida usuário ativo
            norm_email = (email or '').strip().lower()
            cur.execute("SELECT email, ativo FROM usuario WHERE LOWER(email) = %s", (norm_email,))
            urow = cur.fetchone()
            if not urow:
                flash('E-mail não cadastrado.', 'error')
                cur.close(); conn.close()
                return redirect(url_for('esqueci_senha'))
            ativo_val = urow.get('ativo')
            is_active = bool(ativo_val) if ativo_val is not None else True
            if not is_active:
                flash('Usuário inativo. Entre em contato com o administrador.', 'error')
                cur.close(); conn.close()
                return redirect(url_for('esqueci_senha'))
            canonical_email = urow['email']

            # Valida token vinculado ao e-mail canônico
            cur.execute("SELECT * FROM esqueci_senha WHERE email = %s AND token = %s AND status = 'Ativo'", (canonical_email, token))
            req = cur.fetchone()
            if not req:
                flash('Token inválido ou expirado.', 'error')
                cur.close(); conn.close()
                return redirect(url_for('esqueci_senha'))
            # Validade configurável
            requested_at = req['data_solicitacao']
            if isinstance(requested_at, datetime):
                age_seconds = (datetime.utcnow() - requested_at).total_seconds()
                if age_seconds > RESET_TOKEN_EXP_SECONDS:
                    cur2 = conn.cursor()
                    cur2.execute("UPDATE esqueci_senha SET status = 'Expirado' WHERE id_solicitacao = %s", (req['id_solicitacao'],))
                    conn.commit()
                    cur2.close()
                    cur.close(); conn.close()
                    flash('Token expirado. Solicite um novo.', 'error')
                    return redirect(url_for('esqueci_senha'))
            # Atualiza senha do usuário usando e-mail canônico
            nova_hash = generate_password_hash(nova)
            cur2 = conn.cursor()
            cur2.execute("UPDATE usuario SET senha = %s WHERE email = %s", (nova_hash, canonical_email))
            # Expira e remove o token após redefinir
            cur2.execute("UPDATE esqueci_senha SET status = 'Expirado' WHERE id_solicitacao = %s", (req['id_solicitacao'],))
            cur2.execute("DELETE FROM esqueci_senha WHERE id_solicitacao = %s", (req['id_solicitacao'],))
            conn.commit()
            cur2.close()
            cur.close()
            conn.close()
            flash('Senha redefinida com sucesso. Faça login.', 'success')
            audit_log('reset_ok', {'email': canonical_email})
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'ok': True})
            return redirect(url_for('login'))
        except Exception as e:
            try:
                conn.close()
            except Exception:
                pass
            flash(f'Erro ao redefinir senha: {e}', 'error')
            audit_log('reset_error', {'error': str(e)})
            return redirect(url_for('esqueci_senha'))

    # GET
    if not token or not email:
        flash('Token ou e-mail ausente.', 'error')
        return redirect(url_for('esqueci_senha'))
    return render_template('resetar_senha.html', token=token, email=email)

# Validação de código (AJAX)
@app.route('/api/reset/validate', methods=['POST'])
def api_reset_validate():
    email = (request.form.get('email') or '').strip()
    token = (request.form.get('token') or '').strip()
    if not email or not token:
        return jsonify({'ok': False, 'error': 'Informe e-mail e código.'}), 400
    conn = get_db_connection()
    if not conn:
        return jsonify({'ok': False, 'error': 'Falha ao conectar ao banco.'}), 500
    try:
        cur = conn.cursor(row_factory=dict_row)
        # Resolve e-mail canônico do usuário e valida se está ativo
        norm_email = email.lower()
        cur.execute("SELECT email, ativo FROM usuario WHERE LOWER(email) = %s", (norm_email,))
        urow = cur.fetchone()
        if not urow:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'E-mail não cadastrado.'}), 404
        ativo_val = urow.get('ativo')
        is_active = bool(ativo_val) if ativo_val is not None else True
        if not is_active:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Usuário inativo.'}), 403
        canonical_email = urow['email']

        # Valida o token vinculado ao e-mail canônico
        cur.execute("SELECT * FROM esqueci_senha WHERE email = %s AND token = %s AND status = 'Ativo'", (canonical_email, token))
        req = cur.fetchone()
        if not req:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Código inválido ou expirado.'}), 404
        requested_at = req['data_solicitacao']
        if isinstance(requested_at, datetime):
            age_seconds = (datetime.utcnow() - requested_at).total_seconds()
            if age_seconds > RESET_TOKEN_EXP_SECONDS:
                cur2 = conn.cursor()
                cur2.execute("UPDATE esqueci_senha SET status = 'Expirado' WHERE id_solicitacao = %s", (req['id_solicitacao'],))
                conn.commit()
                cur2.close()
                cur.close(); conn.close()
                return jsonify({'ok': False, 'error': 'Código expirado. Solicite novamente.'}), 410
        cur.close(); conn.close()
        return jsonify({'ok': True})
    except Exception as e:
        try: conn.close()
        except Exception: pass
        return jsonify({'ok': False, 'error': f'Erro ao validar código: {e}'}), 500

# Alterar senha via API (AJAX)
@app.route('/api/reset/change', methods=['POST'])
def api_reset_change():
    email = (request.form.get('email') or '').strip()
    token = (request.form.get('token') or '').strip()
    new_password = (request.form.get('new_password') or request.form.get('nova_senha') or '').strip()
    confirm = (request.form.get('confirm_password') or request.form.get('confirmar_senha') or '').strip()
    if not email or not token:
        return jsonify({'ok': False, 'error': 'Token ou e-mail ausente.'}), 400
    if not new_password:
        return jsonify({'ok': False, 'error': 'Informe a nova senha.'}), 400
    if confirm and new_password != confirm:
        return jsonify({'ok': False, 'error': 'As senhas devem coincidir.'}), 400
    if not password_policy_ok(new_password):
        return jsonify({'ok': False, 'error': 'A nova senha deve ter exatamente 8 caracteres, com maiúscula, minúscula, número e símbolo.'}), 400
    conn = get_db_connection()
    if not conn:
        return jsonify({'ok': False, 'error': 'Falha ao conectar ao banco.'}), 500
    try:
        cur = conn.cursor(row_factory=dict_row)
        # Resolve e-mail canônico e valida usuário ativo
        norm_email = email.lower()
        cur.execute("SELECT email, ativo FROM usuario WHERE LOWER(email) = %s", (norm_email,))
        urow = cur.fetchone()
        if not urow:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'E-mail não cadastrado.'}), 404
        ativo_val = urow.get('ativo')
        is_active = bool(ativo_val) if ativo_val is not None else True
        if not is_active:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Usuário inativo.'}), 403
        canonical_email = urow['email']

        # Valida token vinculado ao e-mail canônico
        cur.execute("SELECT * FROM esqueci_senha WHERE email = %s AND token = %s AND status = 'Ativo'", (canonical_email, token))
        req = cur.fetchone()
        if not req:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Token inválido ou expirado.'}), 400
        # Checa expiração
        requested_at = req['data_solicitacao']
        if isinstance(requested_at, datetime):
            age_seconds = (datetime.utcnow() - requested_at).total_seconds()
            if age_seconds > RESET_TOKEN_EXP_SECONDS:
                cur2 = conn.cursor()
                cur2.execute("UPDATE esqueci_senha SET status = 'Expirado' WHERE id_solicitacao = %s", (req['id_solicitacao'],))
                conn.commit()
                cur2.close()
                cur.close(); conn.close()
                return jsonify({'ok': False, 'error': 'Token expirado. Solicite um novo.'}), 400

        # Atualiza senha e expira token
        nova_hash = generate_password_hash(new_password)
        cur2 = conn.cursor()
        cur2.execute("UPDATE usuario SET senha = %s WHERE email = %s", (nova_hash, canonical_email))
        cur2.execute("UPDATE esqueci_senha SET status = 'Expirado' WHERE id_solicitacao = %s", (req['id_solicitacao'],))
        cur2.execute("DELETE FROM esqueci_senha WHERE id_solicitacao = %s", (req['id_solicitacao'],))
        conn.commit()
        cur2.close(); cur.close(); conn.close()
        audit_log('reset_ok_api', {'email': canonical_email})
        return jsonify({'ok': True})
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        audit_log('reset_change_error', {'error': str(e)})
        return jsonify({'ok': False, 'error': f'Erro ao redefinir senha: {str(e)}'}), 500

# Rota para a página de cadastro de alunos
@app.route('/cadastro_alunos', methods=['GET', 'POST'])
def cadastro_alunos():
    # Em ambientes sem autenticação, não restringir acesso
    if request.method == 'POST' and request.form:
        action = (request.form.get('action') or request.form.get('acao') or '').strip().lower()
        if action == 'toggle':
            key = f"{request.remote_addr}:cadastro_alunos_toggle"
            if not check_rate_limit(key, limit=20, window=60):
                flash('Muitas tentativas. Tente novamente em instantes.', 'error')
                audit_log('rate_limit', {'route': 'cadastro_alunos_toggle'})
                return redirect(url_for('cadastro_alunos'))
            ensure_usuario_ativo_column()
            id_usuario = request.form.get('id_usuario')
            if not id_usuario:
                flash('Usuário inválido para alternar status.', 'error')
                return redirect(url_for('cadastro_alunos'))
            conn = get_db_connection()
            if conn:
                try:
                    cur = conn.cursor()
                    cur.execute("SELECT ativo FROM usuario WHERE id_usuario = %s", (id_usuario,))
                    row = cur.fetchone()
                    current_active = True
                    if row is not None:
                        val = row[0] if isinstance(row, tuple) else row
                        current_active = bool(val) if val is not None else True
                    new_active = not current_active
                    cur.execute("UPDATE usuario SET ativo = %s WHERE id_usuario = %s", (new_active, id_usuario))
                    conn.commit()
                    flash('Usuário reativado com sucesso!' if new_active else 'Usuário inativado com sucesso!', 'success')
                    audit_log('usuario_toggle', {'id_usuario': id_usuario, 'ativo': new_active})
                    cur.close(); conn.close()
                    return redirect(url_for('cadastro_alunos'))
                except Exception as e:
                    try:
                        conn.close()
                    except Exception:
                        pass
                    flash(f'Erro ao alternar status: {e}', 'error')
                    audit_log('cadastro_aluno_toggle_error', {'error': str(e)})
                    return redirect(url_for('cadastro_alunos'))
        # Rate limit por IP+rota
        key = f"{request.remote_addr}:cadastro_alunos"
        if not check_rate_limit(key, limit=20, window=60):
            flash('Muitas tentativas. Tente novamente em instantes.', 'error')
            audit_log('rate_limit', {'route': 'cadastro_alunos'})
            return redirect(url_for('cadastro_alunos'))

        action = (request.form.get('action') or request.form.get('acao') or '').strip().lower()
        if action == 'update':
            key_upd = f"{request.remote_addr}:cadastro_alunos_update"
            if not check_rate_limit(key_upd, limit=20, window=60):
                flash('Muitas tentativas. Tente novamente em instantes.', 'error')
                audit_log('rate_limit', {'route': 'cadastro_alunos_update'})
                return redirect(url_for('cadastro_alunos'))
            id_usuario = request.form.get('id_usuario')
            nome = (request.form.get('nome') or request.form.get('nome_user') or '').strip()
            email = (request.form.get('email') or request.form.get('email_user') or '').strip()
            cpf = (request.form.get('cpf') or request.form.get('cpf_user') or '').strip()
            tipo_form = (request.form.get('tipo_usuario') or '').strip()
            # Endereço (opcionais)
            cep = (request.form.get('cep') or request.form.get('cep_user') or '').strip()
            logradouro = (request.form.get('logradouro') or '').strip()
            complemento = (request.form.get('complemento') or '').strip()
            bairro = (request.form.get('bairro') or '').strip()
            cidade = (request.form.get('cidade') or '').strip()
            estado = (request.form.get('estado') or '').strip()
            if not id_usuario:
                flash('Usuário inválido para edição.', 'error')
                return redirect(url_for('cadastro_alunos'))
            if not nome or not email or not cpf:
                flash('Por favor, preencha nome, e-mail e CPF.', 'error')
                return redirect(url_for('cadastro_alunos'))
            # Limite de tamanho do nome
            if len(nome) > 150:
                flash('Nome deve ter no máximo 150 caracteres.', 'error')
                return redirect(url_for('cadastro_alunos'))
            if '@' not in email or '.' not in email:
                flash('E-mail inválido.', 'error')
                return redirect(url_for('cadastro_alunos'))
            if not validar_cpf(cpf):
                flash('CPF inválido.', 'error')
                return redirect(url_for('cadastro_alunos'))
            # Tratamento adicional: CPF com formato válido, mas já associado a outro cadastro
            try:
                conn_chk = get_db_connection()
                if conn_chk:
                    cur_chk = conn_chk.cursor()
                    cpf_digits = re.sub(r'[^0-9]', '', cpf)
                    cur_chk.execute("SELECT 1 FROM usuario WHERE regexp_replace(cpf, '[^0-9]', '', 'g') = %s AND id_usuario <> %s", (cpf_digits, id_usuario))
                    if cur_chk.fetchone():
                        cur_chk.close(); conn_chk.close()
                        flash('CPF informado já está cadastrado para outro usuário. Verifique se é o seu CPF.', 'error')
                        return redirect(url_for('cadastro_alunos'))
                    cur_chk.close(); conn_chk.close()
            except Exception:
                pass
            ensure_usuario_endereco_columns()
            conn = get_db_connection()
            if conn:
                try:
                    cur = conn.cursor()
                    # Bloquear edição se usuário estiver inativo
                    try:
                        ensure_usuario_ativo_column()
                        cur.execute("SELECT COALESCE(ativo, TRUE) FROM usuario WHERE id_usuario = %s", (id_usuario,))
                        row_state = cur.fetchone()
                        is_active = True
                        if row_state is not None:
                            val = row_state[0] if isinstance(row_state, tuple) else row_state
                            is_active = bool(val) if val is not None else True
                        if not is_active:
                            flash('Usuário inativo. Reative para editar.', 'warning')
                            audit_log('cadastro_aluno_update_blocked_inativo', {'id_usuario': id_usuario})
                            cur.close(); conn.close()
                            return redirect(url_for('cadastro_alunos'))
                    except Exception:
                        # Em caso de erro ao verificar, seguir para demais validações
                        pass
                    # Verificar e-mail duplicado em outro usuário
                    cur.execute("SELECT 1 FROM usuario WHERE email = %s AND id_usuario <> %s", (email, id_usuario))
                    if cur.fetchone():
                        flash('E-mail já utilizado por outro usuário.', 'error')
                        cur.close(); conn.close()
                        return redirect(url_for('cadastro_alunos'))
                    tipo_db = 'Aluno'
                    if (tipo_form or '').lower() == 'docente':
                        tipo_db = 'Professor'
                    cur.execute(
                        "UPDATE usuario SET nome = %s, email = %s, cpf = %s, tipo = %s, cep = %s, logradouro = %s, complemento = %s, bairro = %s, cidade = %s, estado = %s WHERE id_usuario = %s",
                        (nome, email, cpf, tipo_db, cep or None, logradouro or None, complemento or None, bairro or None, cidade or None, estado or None, id_usuario)
                    )
                    conn.commit()
                    cur.close(); conn.close()
                    flash('Usuário atualizado com sucesso!', 'success')
                    audit_log('cadastro_aluno_update_ok', {'id_usuario': id_usuario})
                    return redirect(url_for('cadastro_alunos'))
                except Exception as e:
                    try: conn.close()
                    except Exception: pass
                    flash(f'Erro ao atualizar: {e}', 'error')
                    audit_log('cadastro_aluno_update_error', {'error': str(e)})
                    return redirect(url_for('cadastro_alunos'))

        # Fluxo de criação (padrão)
        nome = (request.form.get('nome') or request.form.get('nome_user') or '').strip()
        email = (request.form.get('email') or request.form.get('email_user') or '').strip()
        cpf = (request.form.get('cpf') or request.form.get('cpf_user') or '').strip()
        senha = (request.form.get('senha') or '').strip()
        confirmar_senha = (request.form.get('confirmar_senha') or '').strip()
        curso = request.form.get('curso')
        tipo_form = (request.form.get('tipo_usuario') or '').strip()
        captcha = (request.form.get('captcha') or '').strip()
        is_ajax = (request.headers.get('X-Requested-With') == 'XMLHttpRequest')
        # Endereço (opcionais)
        cep = (request.form.get('cep') or request.form.get('cep_user') or '').strip()
        logradouro = (request.form.get('logradouro') or '').strip()
        complemento = (request.form.get('complemento') or '').strip()
        bairro = (request.form.get('bairro') or '').strip()
        cidade = (request.form.get('cidade') or '').strip()
        estado = (request.form.get('estado') or '').strip()

        # Validações básicas (criação)
        if not nome or not email or not cpf or not senha:
            msg = 'Por favor, preencha todos os campos obrigatórios.'
            if is_ajax:
                return jsonify({'ok': False, 'field': 'nome_user', 'message': msg}), 400
            flash(msg, 'error')
            return redirect(url_for('login', register='1', err='required'))
        # Limite de tamanho do nome
        if len(nome) > 150:
            msg = 'Nome deve ter no máximo 150 caracteres.'
            if is_ajax:
                return jsonify({'ok': False, 'field': 'nome_user', 'message': msg}), 400
            flash(msg, 'error')
            return redirect(url_for('login', register='1', err='nome_len'))
        if senha != confirmar_senha:
            msg = 'As senhas não coincidem.'
            if is_ajax:
                return jsonify({'ok': False, 'field': 'confirmar_senha', 'message': msg}), 400
            flash(msg, 'error')
            return redirect(url_for('login', register='1', err='senha_match'))
        if '@' not in email or '.' not in email:
            msg = 'E-mail inválido.'
            if is_ajax:
                return jsonify({'ok': False, 'field': 'email_user', 'message': msg}), 400
            flash(msg, 'error')
            return redirect(url_for('login', register='1', err='email'))
        # Limite de tamanho do e-mail
        if len(email) > 40:
            msg = 'E-mail deve ter no máximo 40 caracteres.'
            if is_ajax:
                return jsonify({'ok': False, 'field': 'email_user', 'message': msg}), 400
            flash(msg, 'error')
            return redirect(url_for('login', register='1', err='email_len'))
        if not validar_cpf(cpf):
            msg = 'CPF inválido.'
            if is_ajax:
                return jsonify({'ok': False, 'field': 'cpf_user', 'message': msg}), 400
            flash(msg, 'error')
            return redirect(url_for('login', register='1', err='cpf'))
        # Captcha simples
        if str(session.get('captcha_answer')) != captcha:
            msg = 'Captcha incorreto.'
            if is_ajax:
                return jsonify({'ok': False, 'field': 'captcha', 'message': msg}), 400
            flash(msg, 'error')
            return redirect(url_for('login', register='1', err='captcha'))
        # Política de senha (entre 8 e 16 c/ maiúscula, minúscula, número e símbolo)
        if not password_policy_ok(senha):
            msg = 'A senha deve ter entre 8 e 16 caracteres, com maiúscula, minúscula, número e símbolo.'
            if is_ajax:
                return jsonify({'ok': False, 'field': 'senha_reg', 'message': msg}), 400
            flash(msg, 'error')
            return redirect(url_for('login', register='1', err='pwd'))

        ensure_usuario_endereco_columns()
        # Tratamento adicional: CPF com formato válido, mas já associado a outro cadastro
        try:
            conn_chk = get_db_connection()
            if conn_chk:
                cur_chk = conn_chk.cursor()
                cpf_digits = re.sub(r'[^0-9]', '', cpf)
                cur_chk.execute("SELECT 1 FROM usuario WHERE regexp_replace(cpf, '[^0-9]', '', 'g') = %s", (cpf_digits,))
                if cur_chk.fetchone():
                    cur_chk.close(); conn_chk.close()
                    msg = 'CPF informado já está cadastrado para outro usuário. Verifique se é o seu CPF.'
                    audit_log('cadastro_aluno_fail', {'motivo': 'cpf_duplicado', 'cpf': cpf_digits})
                    if is_ajax:
                        return jsonify({'ok': False, 'field': 'cpf_user', 'message': msg}), 400
                    flash(msg, 'error')
                    return redirect(url_for('login', register='1', err='cpf_dup'))
                cur_chk.close(); conn_chk.close()
        except Exception:
            pass
        conn = get_db_connection()
        if conn:
            try:
                cur = conn.cursor()
                # Verificar se o email já existe
                cur.execute("SELECT 1 FROM usuario WHERE email = %s", (email,))
                if cur.fetchone():
                    msg = 'Email já cadastrado'
                    audit_log('cadastro_aluno_fail', {'motivo': 'email_duplicado', 'email': email})
                    if is_ajax:
                        cur.close(); conn.close()
                        return jsonify({'ok': False, 'field': 'email_user', 'message': msg}), 400
                    flash(msg, 'error')
                    cur.close(); conn.close()
                    return redirect(url_for('login', register='1', err='email_dup'))
                # Determinar tipo do usuário no banco: cadastro sempre como 'Aluno'
                tipo_db = 'Aluno'
                # Inserir novo usuário
                senha_hash = generate_password_hash(senha)
                cur.execute(
                    "INSERT INTO usuario (nome, email, cpf, senha, tipo, curso_usuario, cep, logradouro, complemento, bairro, cidade, estado) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                    (nome, email, cpf, senha_hash, tipo_db, curso, cep or None, logradouro or None, complemento or None, bairro or None, cidade or None, estado or None)
                )
                conn.commit()
                flash('Usuário cadastrado com sucesso!', 'success')
                cur.close()
                conn.close()
                audit_log('cadastro_aluno_ok', {'email': email})
                if is_ajax:
                    return jsonify({'ok': True, 'redirect': url_for('login')})
                # Após cadastro bem-sucedido, voltar para a tela de login
                return redirect(url_for('login'))
            except Exception as e:
                flash(f'Erro ao cadastrar aluno: {e}', 'error')
                try:
                    conn.close()
                except Exception:
                    pass
                audit_log('cadastro_aluno_error', {'error': str(e)})
            if is_ajax:
                return jsonify({'ok': False, 'message': 'Erro ao cadastrar aluno.', 'field': None}), 500
            return redirect(url_for('login', register='1', err='error'))
    
    # Buscar cursos para o formulário e alunos para listagem
    cursos = []
    usuarios = []
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor(row_factory=dict_row)
            cur.execute("SELECT * FROM curso ORDER BY nome_curso")
            cursos = cur.fetchall()
            cur.close()
            conn.close()
        except Exception as e:
            flash(f'Erro ao buscar cursos: {e}', 'error')

    # Buscar usuários para listagem
    ensure_usuario_ativo_column()
    conn2 = get_db_connection()
    if conn2:
        try:
            cur2 = conn2.cursor(row_factory=dict_row)
            cur2.execute("""
                SELECT id_usuario, nome, email, cpf, tipo, curso_usuario, foto_perfil,
                       cep, logradouro, complemento, bairro, cidade, estado,
                       COALESCE(ativo, TRUE) AS ativo
                FROM usuario
                ORDER BY id_usuario DESC
            """)
            usuarios = cur2.fetchall()
            cur2.close(); conn2.close()
        except Exception as e:
            try:
                conn2.close()
            except Exception:
                pass
            flash(f'Erro ao buscar usuários: {e}', 'error')

    # Captcha pergunta
    a, b = random.randint(1, 9), random.randint(1, 9)
    session['captcha_answer'] = str(a + b)
    captcha_question = f"Quanto é {a} + {b}?"
    return render_template('cadastro_alunos.html', cursos=cursos, usuarios=usuarios, captcha_question=captcha_question)

# Rota para a página de cadastro de cursos
@app.route('/cadastro_curso', methods=['GET', 'POST'])
@login_required
@roles_required(['Administrador'])
def cadastro_curso():
    if request.method == 'POST' and request.form:
        # Rate limit
        key = f"{request.remote_addr}:cadastro_curso"
        if not check_rate_limit(key, limit=20, window=60):
            flash('Muitas tentativas. Tente novamente em instantes.', 'error')
            audit_log('rate_limit', {'route': 'cadastro_curso'})
            return redirect(url_for('cadastro_curso'))
        nome_curso = request.form.get('nome_curso') or request.form.get('nome')
        descricao = request.form.get('descricao')
        codigo = request.form.get('codigo')
        autorizacao = request.form.get('autorizacao') or request.form.get('portaria')
        action = (request.form.get('action') or request.form.get('acao') or '').strip().lower()
        coordenador_id = request.form.get('coordenador')
        if coordenador_id == '':
            coordenador_id = None
        
        conn = get_db_connection()
        if conn:
            try:
                cur = conn.cursor()
                if action == 'toggle':
                    ensure_curso_ativo_column()
                    id_curso = request.form.get('id_curso')
                    if not id_curso:
                        flash('Curso inválido para alternar status.', 'error')
                        cur.close(); conn.close()
                        return redirect(url_for('cadastro_curso'))
                    # Obtém estado atual
                    conn2 = get_db_connection()
                    row = None
                    if conn2:
                        try:
                            cur2 = conn2.cursor()
                            cur2.execute("SELECT ativo FROM curso WHERE id_curso = %s", (id_curso,))
                            row = cur2.fetchone()
                            cur2.close(); conn2.close()
                        except Exception:
                            try:
                                conn2.close()
                            except Exception:
                                pass
                            row = None
                    current_active = True
                    if row is not None:
                        val = row[0] if isinstance(row, tuple) else row
                        current_active = bool(val) if val is not None else True
                    new_active = not current_active
                    cur.execute("UPDATE curso SET ativo = %s WHERE id_curso = %s", (new_active, id_curso))
                    conn.commit()
                    flash('Curso reativado com sucesso!' if new_active else 'Curso inativado com sucesso!', 'success')
                    audit_log('curso_toggle', {'id_curso': id_curso, 'ativo': new_active})
                    cur.close(); conn.close()
                    return redirect(url_for('cadastro_curso'))
                elif action == 'update':
                    id_curso = request.form.get('id_curso')
                    if not id_curso or not nome_curso:
                        flash('Informe nome e selecione o curso para editar.', 'error')
                        cur.close(); conn.close()
                        audit_log('cadastro_curso_fail', {'motivo': 'update_campos_invalidos'})
                        return redirect(url_for('cadastro_curso'))
                    # Bloqueia edição se o curso estiver inativo
                    try:
                        ensure_curso_ativo_column()
                    except Exception:
                        pass
                    ativo_row = None
                    try:
                        cur.execute("SELECT COALESCE(ativo, TRUE) FROM curso WHERE id_curso = %s", (id_curso,))
                        ativo_row = cur.fetchone()
                    except Exception:
                        ativo_row = None
                    if ativo_row is not None:
                        is_active = ativo_row[0] if isinstance(ativo_row, tuple) else bool(ativo_row)
                        if not is_active:
                            flash('Curso inativo. Reative para editar.', 'warning')
                            audit_log('cadastro_curso_update_bloqueado', {'id_curso': id_curso, 'motivo': 'curso_inativo'})
                            cur.close(); conn.close()
                            return redirect(url_for('cadastro_curso'))
                    cur.execute(
                        "UPDATE curso SET nome_curso = %s, descricao_curso = %s, codigo_curso = %s, autorizacao = %s, id_coordenador = %s WHERE id_curso = %s",
                        (nome_curso, descricao, codigo, autorizacao, coordenador_id, id_curso)
                    )
                    conn.commit()
                    flash('Curso atualizado com sucesso!', 'success')
                    cur.close(); conn.close()
                    audit_log('cadastro_curso_update', {'id_curso': id_curso, 'nome_curso': nome_curso})
                    return redirect(url_for('cadastro_curso'))
                else:
                    # Inserir novo curso
                    if nome_curso:
                        cur.execute(
                            "INSERT INTO curso (nome_curso, descricao_curso, codigo_curso, autorizacao, id_coordenador) VALUES (%s, %s, %s, %s, %s)",
                            (nome_curso, descricao, codigo, autorizacao, coordenador_id)
                        )
                        conn.commit()
                        flash('Curso cadastrado com sucesso!', 'success')
                        cur.close(); conn.close()
                        audit_log('cadastro_curso_ok', {'nome_curso': nome_curso})
                        return redirect(url_for('cadastro_curso'))
                    else:
                        flash('Informe ao menos o nome do curso.', 'error')
                        cur.close(); conn.close()
                        audit_log('cadastro_curso_fail', {'motivo': 'nome_vazio'})
            except Exception as e:
                flash(f'Erro ao processar curso: {e}', 'error')
                audit_log('cadastro_curso_error', {'error': str(e)})
    
    # Buscar professores para o formulário
    professores = []
    cursos = []
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor(row_factory=dict_row)
            cur.execute("SELECT * FROM usuario WHERE tipo = 'Professor' ORDER BY nome")
            professores = cur.fetchall()
            # Buscar cursos já cadastrados
            ensure_curso_ativo_column()
            cur.execute(
                """
                SELECT c.id_curso, c.nome_curso, c.descricao_curso, c.codigo_curso, c.autorizacao, c.ativo, c.id_coordenador, u.nome as coordenador
                FROM curso c
                LEFT JOIN usuario u ON c.id_coordenador = u.id_usuario
                ORDER BY c.nome_curso ASC
                """
            )
            cursos = cur.fetchall()
            cur.close()
            conn.close()
        except Exception as e:
            flash(f'Erro ao buscar professores: {e}', 'error')
    return render_template('cadastro_curso.html', professores=professores, cursos=cursos)

# Rota para a página de publicação
@app.route('/publicacao', methods=['GET', 'POST'])
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def publicacao():
    if request.method == 'POST':
        is_ajax = (request.headers.get('X-Requested-With') == 'XMLHttpRequest')
        # Garante que o enum possua rótulo de pendência
        try:
            ensure_status_enum_pending()
        except Exception:
            pass
        # Rate limit
        key = f"{request.remote_addr}:publicacao"
        if not check_rate_limit(key, limit=15, window=60):
            flash('Muitas tentativas. Tente novamente em instantes.', 'error')
            audit_log('rate_limit', {'route': 'publicacao'})
            return redirect(url_for('publicacao'))
        titulo = (request.form.get('titulo') or request.form.get('titulo_conteudo') or '').strip()
        tipo = (request.form.get('tipo') or request.form.get('tipo_publicacao') or '').strip()
        curso_id = request.form.get('curso')  # pode ser None se não houver campo
        orientador_id = (request.form.get('orientador') or '').strip()
        captcha = (request.form.get('captcha') or '').strip()
        arquivo = request.files.get('conteudo')
        termo_file = request.files.get('termo')
        # Captcha
        if str(session.get('captcha_answer')) != captcha:
            if is_ajax:
                return jsonify({'ok': False, 'message': 'Captcha incorreto.', 'field': 'captcha'}), 400
            flash('Captcha incorreto.', 'error')
            return redirect(url_for('publicacao'))
        if not titulo:
            if is_ajax:
                return jsonify({'ok': False, 'message': 'Informe o título da publicação.', 'field': 'titulo_conteudo'}), 400
            flash('Informe o título da publicação.', 'error')
            return redirect(url_for('publicacao'))
        if not tipo:
            if is_ajax:
                return jsonify({'ok': False, 'message': 'Informe o tipo da publicação.', 'field': 'tipo_publicacao'}), 400
            flash('Informe o tipo da publicação.', 'error')
            return redirect(url_for('publicacao'))
        # Orientador obrigatório e deve ser Professor
        if not orientador_id:
            if is_ajax:
                return jsonify({'ok': False, 'message': 'Selecione o orientador (perfil Professor) para a publicação.', 'field': 'orientador'}), 400
            flash('Selecione o orientador (perfil Professor) para a publicação.', 'error')
            return redirect(url_for('publicacao'))
        if not (arquivo and arquivo.filename):
            if is_ajax:
                return jsonify({'ok': False, 'message': 'Anexe o arquivo de conteúdo para publicar.', 'field': 'conteudo'}), 400
            flash('Anexe o arquivo de conteúdo para publicar.', 'error')
            return redirect(url_for('publicacao'))
        # Termo de autorização obrigatório
        if not (termo_file and termo_file.filename):
            if is_ajax:
                return jsonify({'ok': False, 'message': 'Anexe o termo de autorização.', 'field': 'termo'}), 400
            flash('Anexe o termo de autorização.', 'error')
            return redirect(url_for('publicacao'))
        # Validação de tipos de arquivo
        ALLOW_EXT = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv', '.txt', '.png', '.jpg', '.jpeg', '.webp'}
        ext = os.path.splitext(arquivo.filename)[1].lower()
        if ext not in ALLOW_EXT:
            if is_ajax:
                return jsonify({'ok': False, 'message': 'Tipo de arquivo não permitido.', 'field': 'conteudo'}), 400
            flash('Tipo de arquivo não permitido.', 'error')
            return redirect(url_for('publicacao'))
        ext_termo = os.path.splitext(termo_file.filename)[1].lower()
        if ext_termo not in ALLOW_EXT:
            if is_ajax:
                return jsonify({'ok': False, 'message': 'Tipo de arquivo do termo não permitido.', 'field': 'termo'}), 400
            flash('Tipo de arquivo do termo não permitido.', 'error')
            return redirect(url_for('publicacao'))

        # Gerar nome seguro para o arquivo
        filename = secure_filename(arquivo.filename)
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        novo_filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], novo_filename)
        
        # Salvar os arquivos
        arquivo.save(filepath)
        # salva termo de autorização também (não referenciado na tabela)
        try:
            term_name = secure_filename(termo_file.filename)
            ts_term = datetime.now().strftime('%Y%m%d%H%M%S')
            new_term_name = f"{ts_term}_termo_{term_name}"
            termpath = os.path.join(app.config['UPLOAD_FOLDER'], new_term_name)
            termo_file.save(termpath)
        except Exception:
            pass
        
        conn = get_db_connection()
        if conn:
            try:
                cur = conn.cursor()
                # valida orientador como Professor
                try:
                    cur_prof = conn.cursor()
                    cur_prof.execute("SELECT 1 FROM usuario WHERE id_usuario = %s AND tipo::text IN ('Professor','Docente')", (orientador_id,))
                    if not cur_prof.fetchone():
                        cur_prof.close()
                        flash('Orientador inválido: selecione um usuário com perfil Professor.', 'error')
                        try:
                            conn.close()
                        except Exception:
                            pass
                        return redirect(url_for('publicacao'))
                    cur_prof.close()
                except Exception:
                    # Se falhar a validação, impedir inserção
                    flash('Falha ao validar o orientador. Tente novamente.', 'error')
                    try:
                        conn.close()
                    except Exception:
                        pass
                    return redirect(url_for('publicacao'))
                # Define status inicial: Aluno -> Pendente; Docente/Admin -> Publicado
                try:
                    user_tipo = (session.get('user_tipo') or session.get('role') or '').strip()
                except Exception:
                    user_tipo = ''
                initial_status = status_label('Pendente') if user_tipo == 'Aluno' else status_label('Publicado')

                # Inserir nova publicação (id_autor e id_curso podem ser None)
                cur.execute(
                    """INSERT INTO publicacao 
                       (titulo, data_publicacao, id_autor, id_curso, tipo, status, arquivo, nome_arquivo, assuntos_relacionados, data_autoria, id_orientador) 
                       VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                    (titulo, datetime.now(), session.get('user_id'), curso_id, tipo or '', initial_status, 
                     filepath, novo_filename, None, None, orientador_id)
                )
                conn.commit()
                try:
                    pend_lbl = status_label('Pendente')
                except Exception:
                    pend_lbl = 'Pendente'
                if is_ajax:
                    msg = 'Submissão criada e aguardando avaliação do orientador.' if _norm(initial_status) == _norm(pend_lbl) else 'Publicação criada e publicada.'
                    cur.close(); conn.close()
                    audit_log('publicacao_ok', {'titulo': titulo, 'arquivo': novo_filename})
                    return jsonify({'ok': True, 'message': msg})
                if _norm(initial_status) == _norm(pend_lbl):
                    flash('Submissão criada e aguardando avaliação do orientador.', 'success')
                else:
                    flash('Publicação criada e publicada.', 'success')
                cur.close()
                conn.close()
                audit_log('publicacao_ok', {'titulo': titulo, 'arquivo': novo_filename})
            except Exception as e:
                if is_ajax:
                    try:
                        conn.close()
                    except Exception:
                        pass
                    audit_log('publicacao_error', {'error': str(e)})
                    return jsonify({'ok': False, 'message': f'Erro ao publicar: {e}', 'field': None}), 500
                flash(f'Erro ao publicar: {e}', 'error')
                try:
                    conn.close()
                except Exception:
                    pass
                audit_log('publicacao_error', {'error': str(e)})
        if is_ajax:
            return jsonify({'ok': True, 'message': 'Publicação criada.'})
        return redirect(url_for('publicacao'))

    # Buscar cursos e tipos de publicação para o formulário e listar últimas publicações
    cursos = []
    tipos = []
    professores = []
    publicacoes = []
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor(row_factory=dict_row)
            # Corrige publicações sem avaliação marcadas como Publicada
            # Executa em conexão isolada para evitar abortar a transação usada nas consultas abaixo
            try:
                pend_label = status_label('Pendente')
                pub_label = status_label('Publicado')
                conn_fix = get_db_connection()
                if conn_fix:
                    with conn_fix.cursor() as cur_fix:
                        cur_fix.execute(
                            """
                            UPDATE publicacao p
                            SET status = %s
                            WHERE p.status = %s
                              AND NOT EXISTS (
                                SELECT 1 FROM avaliacao a WHERE a.id_publicacao = p.id_publicacao
                              )
                            """,
                            (pend_label, pub_label)
                        )
                        conn_fix.commit()
                    conn_fix.close()
            except Exception:
                try:
                    conn_fix.close()
                except Exception:
                    pass
            # Garante colunas opcionais e filtra somente cursos ativos (ou sem coluna ativo)
            try:
                ensure_curso_ativo_column()
            except Exception:
                pass
            cur.execute("""
                SELECT id_curso, nome_curso
                FROM curso
                WHERE COALESCE(ativo, TRUE) = TRUE
                ORDER BY nome_curso
            """)
            cursos = cur.fetchall()
            
            cur.execute("SELECT * FROM tipos_de_publicacao ORDER BY nome_tipo")
            tipos = cur.fetchall()

            # Filtra orientadores/professores ativos
            try:
                ensure_usuario_ativo_column()
            except Exception:
                pass
            cur.execute("""
                SELECT id_usuario, nome
                FROM usuario
                WHERE tipo::text IN ('Professor','Docente')
                  AND COALESCE(ativo, TRUE) = TRUE
                ORDER BY nome
            """)
            professores = cur.fetchall()

            # Filtragem de "Últimas publicações" conforme perfil
            user_role = (session.get('user_tipo') or session.get('role') or '').strip()
            uid = session.get('user_id')
            if user_role == 'Aluno' and uid:
                # Aluno: ver todas Publicadas; ver também Pendente/Indeferida apenas se forem do próprio usuário
                try:
                    pub_label = status_label('Publicado')
                except Exception:
                    pub_label = 'Publicado'
                try:
                    pend_label = status_label('Pendente')
                except Exception:
                    pend_label = 'Pendente'
                # Aceita tanto "Reprovado" quanto "Indeferido" conforme enum
                try:
                    indefer_label = status_label('Indeferido')
                except Exception:
                    try:
                        indefer_label = status_label('Reprovado')
                    except Exception:
                        indefer_label = 'Indeferido'

                cur.execute(
                    """
                    SELECT 
                      p.id_publicacao,
                      p.titulo,
                      p.tipo,
                      c.nome_curso AS curso,
                      u.nome AS autor_nome,
                      p.nome_arquivo,
                      p.data_publicacao,
                      p.status
                    FROM publicacao p
                    LEFT JOIN curso c ON c.id_curso = p.id_curso
                    LEFT JOIN usuario u ON u.id_usuario = p.id_autor
                    WHERE p.status = %s
                       OR (p.id_autor = %s AND p.status IN (%s, %s))
                    ORDER BY p.titulo ASC, p.data_publicacao DESC, p.id_publicacao DESC
                    LIMIT 20
                    """,
                    (pub_label, uid, pend_label, indefer_label)
                )
            else:
                # Administrador/Docente: mantém visão completa
                cur.execute(
                    """
                    SELECT 
                      p.id_publicacao,
                      p.titulo,
                      p.tipo,
                      c.nome_curso AS curso,
                      u.nome AS autor_nome,
                      p.nome_arquivo,
                      p.data_publicacao,
                      p.status
                    FROM publicacao p
                    LEFT JOIN curso c ON c.id_curso = p.id_curso
                    LEFT JOIN usuario u ON u.id_usuario = p.id_autor
                    ORDER BY p.titulo ASC, p.data_publicacao DESC, p.id_publicacao DESC
                    LIMIT 20
                    """
                )
            publicacoes = cur.fetchall()
            
            cur.close()
            conn.close()
        except Exception as e:
            # Não exibir erros na tela durante o carregamento inicial (GET)
            # Registrar em log para diagnóstico sem interromper a experiência do usuário
            try:
                audit_log('vinculacao_fetch_error', {'error': str(e)})
            except Exception:
                pass
    # Captcha pergunta
    a, b = random.randint(1, 9), random.randint(1, 9)
    session['captcha_answer'] = str(a + b)
    captcha_question = f"Quanto é {a} + {b}?"
    return render_template('publicacao.html', cursos=cursos, tipos=tipos, professores=professores, publicacoes=publicacoes, captcha_question=captcha_question)

# Rota de download da publicação com nome do arquivo igual ao título
@app.route('/download_publicacao/<int:id_publicacao>')
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def download_publicacao(id_publicacao):
    conn = get_db_connection()
    if not conn:
        flash('Falha ao obter conexão para download.', 'error')
        return redirect(url_for('publicacao'))
    try:
        cur = conn.cursor(row_factory=dict_row)
        cur.execute("""
            SELECT nome_arquivo, titulo, status
            FROM publicacao
            WHERE id_publicacao = %s
            LIMIT 1
        """, (id_publicacao,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        if not row or not row.get('nome_arquivo'):
            flash('Publicação não encontrada ou sem arquivo.', 'error')
            return redirect(url_for('publicacao'))
        # Bloqueia download se não estiver Publicada
        try:
            pub_label = status_label('Publicado')
        except Exception:
            pub_label = 'Publicado'
        if _norm(str(row.get('status') or '')) != _norm(pub_label):
            return make_response('Download indisponível: publicação pendente ou não publicada.', 403)
        stored_name = row['nome_arquivo']
        titulo = (row['titulo'] or 'publicacao').strip()
        upload_dir = app.config['UPLOAD_FOLDER']
        # Verifica se o arquivo existe fisicamente
        full_path = os.path.join(upload_dir, stored_name)
        if not os.path.exists(full_path):
            # Arquivo ausente: sempre retorna um PDF explicando o problema
            preview_dir = ensure_previews_dir()
            preview_name = f'preview_pub_{id_publicacao}.pdf'
            preview_path = os.path.join(preview_dir, preview_name)
            safe_title = secure_filename(titulo) or 'publicacao'
            download_name = f"{safe_title}.pdf"
            try:
                make_error_pdf(preview_path, 'Arquivo não encontrado', f'O arquivo da publicação (id {id_publicacao}) não está disponível no servidor.')
                resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                try:
                    resp.headers['Content-Length'] = os.path.getsize(preview_path)
                except Exception:
                    pass
                return resp
            except Exception:
                # Último recurso: PDF mínimo
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                doc = SimpleDocTemplate(preview_path, pagesize=A4)
                doc.build([Paragraph('Arquivo da publicação não encontrado.', getSampleStyleSheet()['Normal'])])
                resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                return resp
        # Preserva a extensão original para evitar problemas ao abrir o arquivo
        ext = os.path.splitext(stored_name)[1]
        safe_title = secure_filename(titulo) or 'publicacao'
        download_name = f"{safe_title}{ext}"
        resp = send_from_directory(upload_dir, stored_name, as_attachment=True, download_name=download_name)
        # Define explicitamente Content-Length para permitir barra de progresso no front
        try:
            size_bytes = os.path.getsize(full_path)
            resp.headers['Content-Length'] = size_bytes
        except Exception:
            size_bytes = None
            pass
        try:
            ctype = mimetypes.guess_type(full_path)[0] or 'application/octet-stream'
        except Exception:
            ctype = 'application/octet-stream'
        # Auditoria de download: quem (via sessão), quando (timestamp no audit_log) e qual arquivo
        audit_log('download_publicacao', {
            'id_publicacao': id_publicacao,
            'arquivo': stored_name,
            'nome_download': download_name,
            'size_bytes': size_bytes,
            'content_type': ctype
        })
        return resp
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        flash(f'Erro ao preparar download: {e}', 'error')
        try:
            audit_log('download_publicacao_error', {
                'id_publicacao': id_publicacao,
                'error': str(e)
            })
        except Exception:
            pass
        return redirect(url_for('publicacao'))

# Rota de pré-visualização de publicação para formatos Office
@app.route('/preview_publicacao/<int:id_publicacao>')
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def preview_publicacao(id_publicacao):
    from html import escape
    try:
        conn = get_db_connection()
        cur = conn.cursor(row_factory=dict_row)
        cur.execute("""
            SELECT titulo, nome_arquivo, arquivo
            FROM publicacao
            WHERE id_publicacao = %s
            LIMIT 1
        """, (id_publicacao,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        if not row or not row.get('nome_arquivo'):
            return make_response('<div style="padding:12px;color:#dc2626;">Publicação não encontrada ou sem arquivo.</div>', 404)
        titulo = (row.get('titulo') or '').strip()
        stored_name = row['nome_arquivo']
        upload_dir = app.config['UPLOAD_FOLDER']
        full_path = os.path.join(upload_dir, stored_name)
        if not os.path.exists(full_path):
            return make_response('<div style="padding:12px;color:#dc2626;">Arquivo não encontrado no servidor.</div>', 404)
        ext = os.path.splitext(stored_name)[1].lower()

        html_content = ''
        if ext == '.docx':
            try:
                from docx import Document
                doc = Document(full_path)
                parts = []
                parts.append('<div style="font-family: ui-sans-serif, system-ui; color:#1f2937;">')
                parts.append(f'<h3 style="margin:0 0 8px 0; font-weight:600;">{escape(titulo)}</h3>')
                count = 0
                for p in doc.paragraphs:
                    text = p.text.strip()
                    if text:
                        parts.append(f'<p style="margin:6px 0;">{escape(text)}</p>')
                        count += 1
                        if count >= 120:
                            parts.append('<p style="color:#6b7280;">Pré-visualização truncada…</p>')
                            break
                parts.append('</div>')
                html_content = ''.join(parts)
            except Exception as e:
                html_content = f'<div style="padding:12px;color:#dc2626;">Falha ao gerar pré-visualização DOCX: {escape(str(e))}</div>'

        elif ext in ('.xlsx',):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(full_path, read_only=True, data_only=True)
                ws = wb.active
                parts = []
                parts.append('<div style="font-family: ui-sans-serif, system-ui; color:#1f2937;">')
                parts.append(f'<h3 style="margin:0 0 8px 0; font-weight:600;">{escape(titulo)}</h3>')
                parts.append('<div style="overflow:auto; border:1px solid #e5e7eb; border-radius:6px;">')
                parts.append('<table style="border-collapse:collapse; width:100%;">')
                max_rows = 50
                max_cols = 20
                for row_cells in ws.iter_rows(min_row=1, max_row=max_rows, min_col=1, max_col=max_cols):
                    parts.append('<tr>')
                    for cell in row_cells:
                        val = cell.value
                        txt = '' if val is None else escape(str(val))
                        parts.append(f'<td style="border:1px solid #e5e7eb; padding:6px; font-size:14px;">{txt}</td>')
                    parts.append('</tr>')
                parts.append('</table></div></div>')
                html_content = ''.join(parts)
            except Exception as e:
                html_content = f'<div style="padding:12px;color:#dc2626;">Falha ao gerar pré-visualização XLSX: {escape(str(e))}</div>'

        elif ext in ('.xls',):
            try:
                import xlrd
                book = xlrd.open_workbook(full_path)
                sheet = book.sheet_by_index(0)
                parts = []
                parts.append('<div style="font-family: ui-sans-serif, system-ui; color:#1f2937;">')
                parts.append(f'<h3 style="margin:0 0 8px 0; font-weight:600;">{escape(titulo)}</h3>')
                parts.append('<div style="overflow:auto; border:1px solid #e5e7eb; border-radius:6px;">')
                parts.append('<table style="border-collapse:collapse; width:100%;">')
                max_rows = min(50, sheet.nrows)
                max_cols = min(20, sheet.ncols)
                for r in range(max_rows):
                    parts.append('<tr>')
                    for c in range(max_cols):
                        val = sheet.cell_value(r, c)
                        txt = '' if val is None else escape(str(val))
                        parts.append(f'<td style="border:1px solid #e5e7eb; padding:6px; font-size:14px;">{txt}</td>')
                    parts.append('</tr>')
                parts.append('</table></div></div>')
                html_content = ''.join(parts)
            except Exception as e:
                html_content = f'<div style="padding:12px;color:#dc2626;">Falha ao gerar pré-visualização XLS: {escape(str(e))}</div>'
        else:
            return make_response('<div style="padding:12px;color:#6b7280;">Pré-visualização não suportada por esta rota.</div>', 400)

        resp = make_response(html_content)
        resp.headers['Content-Type'] = 'text/html; charset=utf-8'
        resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        return resp
    except Exception as e:
        from html import escape as esc
        return make_response(f'<div style="padding:12px;color:#dc2626;">Erro ao gerar pré-visualização: {esc(str(e))}</div>', 500)

# Utilitários para conversão automática para PDF

def ensure_previews_dir() -> str:
    d = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'previews')
    os.makedirs(d, exist_ok=True)
    return d


def try_libreoffice_convert(input_path: str, outdir: str):
    """Tenta converter via LibreOffice (soffice). Retorna (ok, caminho_pdf)."""
    try:
        import shutil, subprocess
        soffice = shutil.which('soffice')
        if not soffice:
            return (False, None)
        res = subprocess.run(
            [soffice, '--headless', '--convert-to', 'pdf', '--outdir', outdir, input_path],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60
        )
        if res.returncode == 0:
            base = os.path.splitext(os.path.basename(input_path))[0]
            pdf_path = os.path.join(outdir, base + '.pdf')
            if os.path.exists(pdf_path):
                return (True, pdf_path)
        return (False, None)
    except Exception:
        return (False, None)


def docx_to_pdf_reportlab(input_path: str, out_pdf_path: str):
    """Fallback robusto DOCX→PDF usando ReportLab (texto e tabelas)."""
    from docx import Document
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    normal = styles['Normal']
    title_style = ParagraphStyle('title', parent=styles['Heading2'])

    def safe_text(s: str) -> str:
        try:
            return (s or '').replace('\xa0', ' ').replace('\u200b', ' ').strip()
        except Exception:
            try:
                return str(s).strip()
            except Exception:
                return ''

    story = []

    # Tenta abrir o DOCX; se falhar, gera nota em PDF
    try:
        doc = Document(input_path)
    except Exception:
        doc = None
        story.append(Paragraph('Falha ao abrir o arquivo DOCX.', title_style))
        story.append(Spacer(1, 12))

    if doc:
        # Parágrafos
        for p in doc.paragraphs:
            text = safe_text(p.text)
            if text:
                try:
                    story.append(Paragraph(text, normal))
                    story.append(Spacer(1, 6))
                except Exception:
                    # Se houver caracteres não suportados, remove-os
                    text_ascii = text.encode('ascii', 'ignore').decode('ascii')
                    story.append(Paragraph(text_ascii, normal))
                    story.append(Spacer(1, 6))
        # Tabelas (conteúdo textual)
        try:
            for t in getattr(doc, 'tables', []):
                data = []
                for row in t.rows:
                    cells = []
                    for cell in row.cells:
                        cells.append(safe_text(cell.text))
                    data.append(cells)
                if data:
                    tbl = Table(data)
                    tbl.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#94a3b8')),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e2e8f0')),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT')
                    ]))
                    story.append(tbl)
                    story.append(Spacer(1, 8))
        except Exception:
            pass

    if not story:
        story.append(Paragraph('Documento DOCX sem conteúdo textual suportado.', normal))

    SimpleDocTemplate(
        out_pdf_path,
        pagesize=A4,
        rightMargin=24,
        leftMargin=24,
        topMargin=24,
        bottomMargin=24
    ).build(story)


def excel_to_pdf_reportlab(input_path: str, out_pdf_path: str):
    """Fallback simples Excel→PDF mostrando até 50 linhas e 20 colunas."""
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    data = []
    styles = getSampleStyleSheet()
    ext = os.path.splitext(input_path)[1].lower()
    if ext == '.xlsx':
        import openpyxl
        wb = openpyxl.load_workbook(input_path, read_only=True, data_only=True)
        ws = wb.active
        max_rows = 50
        max_cols = 20
        for r in ws.iter_rows(min_row=1, max_row=max_rows, min_col=1, max_col=max_cols):
            row = []
            for c in r:
                val = c.value
                row.append('' if val is None else str(val))
            data.append(row)
    else:
        import xlrd
        book = xlrd.open_workbook(input_path)
        sheet = book.sheet_by_index(0)
        max_rows = min(50, sheet.nrows)
        max_cols = min(20, sheet.ncols)
        for rr in range(max_rows):
            row = []
            for cc in range(max_cols):
                val = sheet.cell_value(rr, cc)
                row.append('' if val is None else str(val))
            data.append(row)
    doc = SimpleDocTemplate(out_pdf_path, pagesize=A4)
    story = []
    if data:
        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.25, colors.lightgrey),
            ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
            ('FONTSIZE', (0,0), (-1,-1), 10),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ROWHEIGHT', (0,0), (-1,-1), 16),
        ]))
        story.append(Paragraph('Pré-visualização de planilha (máx. 50 linhas, 20 colunas)', styles['Italic']))
        story.append(Spacer(1, 8))
        story.append(t)
    else:
        story.append(Paragraph('Sem dados para exibir.', styles['Normal']))
    doc.build(story)

# ---------- Fallbacks adicionais: imagem, texto e CSV em PDF ----------

def image_to_pdf_reportlab(input_path: str, out_pdf_path: str):
    """Converte PNG/JPG para PDF em uma página."""
    from reportlab.platypus import SimpleDocTemplate, Image, Spacer
    from reportlab.lib.pagesizes import A4
    doc = SimpleDocTemplate(out_pdf_path, pagesize=A4, rightMargin=24, leftMargin=24, topMargin=24, bottomMargin=24)
    page_width = A4[0] - 48
    img = Image(input_path, width=page_width, height=None)
    story = [img, Spacer(1, 6)]
    doc.build(story)


def text_to_pdf_reportlab(input_path: str, out_pdf_path: str):
    """Converte TXT em PDF como texto simples."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import A4
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(out_pdf_path, pagesize=A4, rightMargin=24, leftMargin=24, topMargin=24, bottomMargin=24)
    story = []
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except Exception:
        with open(input_path, 'r', encoding='latin-1') as f:
            lines = f.readlines()
    for ln in lines[:2000]:
        text = (ln or '').rstrip('\n')
        if text:
            story.append(Paragraph(text, styles['Normal']))
            story.append(Spacer(1, 6))
    doc.build(story)


def csv_to_pdf_reportlab(input_path: str, out_pdf_path: str):
    """Converte CSV simples em PDF com tabela."""
    import csv
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    styles = getSampleStyleSheet()
    rows = []
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            sample = f.read(1024)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample)
            except Exception:
                dialect = csv.excel
            reader = csv.reader(f, dialect)
            for r in reader:
                rows.append([str(c) for c in r])
    except Exception:
        with open(input_path, 'r', encoding='latin-1') as f:
            reader = csv.reader(f)
            for r in reader:
                rows.append([str(c) for c in r])
    doc = SimpleDocTemplate(out_pdf_path, pagesize=A4, rightMargin=24, leftMargin=24, topMargin=24, bottomMargin=24)
    story = []
    if rows:
        t = Table(rows[:200])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e2e8f0')),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#94a3b8')),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT')
        ]))
        story.append(Paragraph('CSV (preview)', styles['Heading5']))
        story.append(Spacer(1, 8))
        story.append(t)
    else:
        story.append(Paragraph('CSV vazio.', styles['Normal']))
    doc.build(story)


def make_error_pdf(out_pdf_path: str, title: str, message: str):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('title', parent=styles['Heading2'], textColor=colors.HexColor('#dc2626'))
    doc = SimpleDocTemplate(out_pdf_path, pagesize=A4, rightMargin=24, leftMargin=24, topMargin=24, bottomMargin=24)
    story = [Paragraph(title, title_style), Spacer(1, 12), Paragraph(message, styles['Normal'])]
    try:
        doc.build(story)
    except Exception:
        # Último recurso: tenta texto simples
        doc = SimpleDocTemplate(out_pdf_path, pagesize=A4)
        doc.build([Paragraph(message, styles['Normal'])])

# Rota de pré-visualização com PDF automático
@app.route('/preview_pdf_publicacao/<int:id_publicacao>')
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def preview_pdf_publicacao(id_publicacao):
    try:
        conn = get_db_connection()
        cur = conn.cursor(row_factory=dict_row)
        cur.execute("""
            SELECT titulo, nome_arquivo, arquivo
            FROM publicacao
            WHERE id_publicacao = %s
            LIMIT 1
        """, (id_publicacao,))
        row = cur.fetchone()
        cur.close()
        conn.close()
        if not row:
            return make_response('<div style="padding:12px;color:#dc2626;">Publicação não encontrada.</div>', 404)
        upload_dir = app.config['UPLOAD_FOLDER']
        stored_name = (row.get('nome_arquivo') or '').strip()
        full_path = os.path.join(upload_dir, stored_name) if stored_name else ''
        # Fallback: usa caminho completo armazenado em 'arquivo' ou procura pelo nome dentro de uploads
        if not full_path or not os.path.exists(full_path):
            alt = (row.get('arquivo') or '').strip()
            if alt and os.path.exists(alt):
                full_path = alt
            elif stored_name:
                # procura por arquivo com o mesmo nome em subpastas de uploads
                try:
                    for root, dirs, files in os.walk(upload_dir):
                        if stored_name in files:
                            full_path = os.path.join(root, stored_name)
                            break
                except Exception:
                    pass
        # Se ainda não encontrado, deixamos seguir para cache/erro

        # Sempre tenta servir um PDF em cache se já existir
        preview_dir = ensure_previews_dir()
        preview_name = f'preview_pub_{id_publicacao}.pdf'
        preview_path = os.path.join(preview_dir, preview_name)
        if os.path.exists(preview_path):
            # Usa cache se estiver mais novo que a fonte; se a fonte não existir, serve cache
            try:
                if (os.path.exists(full_path) and os.path.getmtime(preview_path) >= os.path.getmtime(full_path)):
                    return send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=False)
            except Exception:
                pass
            # Se a fonte não existir, ainda assim entrega o cache
            if not os.path.exists(full_path):
                return send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=False)
            # Cache antigo: força regeneração removendo arquivo
            try:
                os.remove(preview_path)
            except Exception:
                pass

        # Se não houver cache e a fonte não existir, gera PDF de erro
        if not os.path.exists(full_path):
            try:
                make_error_pdf(preview_path, 'Arquivo não encontrado', f'O arquivo da publicação (id {id_publicacao}) não está disponível no servidor.')
                resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=False)
                return resp
            except Exception:
                # PDF mínimo
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                doc = SimpleDocTemplate(preview_path, pagesize=A4)
                doc.build([Paragraph('Arquivo da publicação não encontrado.', getSampleStyleSheet()['Normal'])])
                return send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=False)
        ext = os.path.splitext(stored_name)[1].lower()
        if ext not in ('.doc', '.docx', '.xls', '.xlsx'):
            return make_response('<div style="padding:12px;color:#6b7280;">Formato não suportado para conversão automática.</div>', 400)

        # Cache simples: usa preview se for mais novo que a fonte
        try:
            if os.path.exists(preview_path):
                src_m = os.path.getmtime(full_path)
                dst_m = os.path.getmtime(preview_path)
                if dst_m >= src_m:
                    return send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=False)
        except Exception:
            pass

        # Tenta LibreOffice
        ok, lo_pdf = try_libreoffice_convert(full_path, preview_dir)
        if ok and lo_pdf and os.path.exists(lo_pdf):
            try:
                import shutil
                shutil.copyfile(lo_pdf, preview_path)
            except Exception:
                preview_path = lo_pdf
            return send_from_directory(os.path.dirname(preview_path), os.path.basename(preview_path), mimetype='application/pdf', as_attachment=False)

        # Fallbacks
        try:
            if ext in ('.docx',):
                docx_to_pdf_reportlab(full_path, preview_path)
            elif ext in ('.xlsx', '.xls'):
                excel_to_pdf_reportlab(full_path, preview_path)
            else:
                return make_response('<div style="padding:12px;color:#6b7280;">Converter este formato requer LibreOffice no servidor.</div>', 400)
            return send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=False)
        except Exception as e:
            from html import escape as esc
            return make_response(f'<div style="padding:12px;color:#dc2626;">Falha ao gerar PDF: {esc(str(e))}</div>', 500)
    except Exception as e:
        from html import escape as esc
        return make_response(f'<div style="padding:12px;color:#dc2626;">Erro ao preparar pré-visualização: {esc(str(e))}</div>', 500)

# Download como PDF (força PDF em vez do arquivo original)
@app.route('/download_pdf_publicacao/<int:id_publicacao>')
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def download_pdf_publicacao(id_publicacao):
    try:
        conn = get_db_connection()
        cur = conn.cursor(row_factory=dict_row)
        cur.execute("""
            SELECT titulo, nome_arquivo, arquivo, status
            FROM publicacao
            WHERE id_publicacao = %s
            LIMIT 1
        """, (id_publicacao,))
        row = cur.fetchone()
        cur.close(); conn.close()
        if not row:
            # gera PDF mínimo informando que não encontrou a publicação
            pdf_bytes = generate_error_pdf(f'Publicação {id_publicacao} não encontrada.')
            return send_file(io.BytesIO(pdf_bytes), mimetype='application/pdf', as_attachment=True, download_name=f'publicacao_{id_publicacao}.pdf')
        # Bloqueia download se não estiver Publicada
        try:
            pub_label = status_label('Publicado')
        except Exception:
            pub_label = 'Publicado'
        if _norm(str(row.get('status') or '')) != _norm(pub_label):
            return make_response('Download indisponível: publicação pendente ou não publicada.', 403)
        titulo = (row.get('titulo') or 'publicacao').strip()
        stored_name = (row.get('nome_arquivo') or '').strip()
        upload_dir = app.config['UPLOAD_FOLDER']
        full_path = os.path.join(upload_dir, stored_name) if stored_name else ''
        # Fallback: tenta usar caminho completo em 'arquivo' ou localizar em subpastas
        if not full_path or not os.path.exists(full_path):
            alt = (row.get('arquivo') or '').strip()
            if alt and os.path.exists(alt):
                full_path = alt
            elif stored_name:
                try:
                    for root, dirs, files in os.walk(upload_dir):
                        if stored_name in files:
                            full_path = os.path.join(root, stored_name)
                            break
                except Exception:
                    pass
        if not os.path.exists(full_path):
            # Arquivo ausente: usa cache de preview se existir; caso contrário, gera PDF de erro
            preview_dir = ensure_previews_dir()
            preview_name = f'preview_pub_{id_publicacao}.pdf'
            preview_path = os.path.join(preview_dir, preview_name)
            safe_title = secure_filename(titulo) or 'publicacao'
            download_name = f"{safe_title}.pdf"
            if os.path.exists(preview_path):
                resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                try:
                    resp.headers['Content-Length'] = os.path.getsize(preview_path)
                except Exception:
                    pass
                return resp
            try:
                make_error_pdf(preview_path, 'Arquivo não encontrado', f'O arquivo da publicação (id {id_publicacao}) não está disponível no servidor.')
                resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                try:
                    resp.headers['Content-Length'] = os.path.getsize(preview_path)
                except Exception:
                    pass
                return resp
            except Exception:
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                doc = SimpleDocTemplate(preview_path, pagesize=A4)
                doc.build([Paragraph('Arquivo da publicação não encontrado.', getSampleStyleSheet()['Normal'])])
                resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                return resp

        preview_dir = ensure_previews_dir()
        preview_name = f'preview_pub_{id_publicacao}.pdf'
        preview_path = os.path.join(preview_dir, preview_name)
        safe_title = secure_filename(titulo) or 'publicacao'
        download_name = f"{safe_title}.pdf"

        # Se já existe PDF em cache e está mais novo que a fonte, usa-o
        if os.path.exists(preview_path):
            try:
                if os.path.getmtime(preview_path) >= os.path.getmtime(full_path):
                    resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                    try:
                        size_bytes = os.path.getsize(preview_path)
                        resp.headers['Content-Length'] = size_bytes
                    except Exception:
                        pass
                    return resp
            except Exception:
                pass
            # Cache desatualizado: remove para forçar regeneração
            try:
                os.remove(preview_path)
            except Exception:
                pass

        # Caso contrário, tenta gerar (mesma lógica do preview)
        ext = os.path.splitext(stored_name)[1].lower()

        # Se já é PDF, apenas força o nome baseado no título
        if ext == '.pdf':
            resp = send_file(full_path, mimetype='application/pdf', as_attachment=True, download_name=download_name)
            try:
                resp.headers['Content-Length'] = os.path.getsize(full_path)
            except Exception:
                pass
            return resp

        # Office: tenta LibreOffice e depois fallbacks internos
        if ext in ('.doc', '.docx', '.xls', '.xlsx'):
            ok, lo_pdf = try_libreoffice_convert(full_path, preview_dir)
            if ok and lo_pdf and os.path.exists(lo_pdf):
                try:
                    import shutil
                    shutil.copyfile(lo_pdf, preview_path)
                except Exception:
                    preview_path = lo_pdf
                resp = send_from_directory(os.path.dirname(preview_path), os.path.basename(preview_path), mimetype='application/pdf', as_attachment=True, download_name=download_name)
                try:
                    resp.headers['Content-Length'] = os.path.getsize(preview_path)
                except Exception:
                    pass
                return resp
            # Fallbacks internos
            try:
                if ext == '.docx':
                    docx_to_pdf_reportlab(full_path, preview_path)
                else:
                    excel_to_pdf_reportlab(full_path, preview_path)
                resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                try:
                    resp.headers['Content-Length'] = os.path.getsize(preview_path)
                except Exception:
                    pass
                return resp
            except Exception as e:
                # Falhou a conversão: gera PDF simplificado para garantir formato .pdf
                try:
                    make_error_pdf(preview_path, 'Conversão não disponível', f'Falha ao converter {stored_name}. PDF simplificado gerado.')
                    resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                    try:
                        resp.headers['Content-Length'] = os.path.getsize(preview_path)
                    except Exception:
                        pass
                    return resp
                except Exception:
                    # Último recurso: gera um PDF mínimo
                    from reportlab.platypus import SimpleDocTemplate, Paragraph
                    from reportlab.lib.pagesizes import A4
                    from reportlab.lib.styles import getSampleStyleSheet
                    doc = SimpleDocTemplate(preview_path, pagesize=A4)
                    doc.build([Paragraph('Conversão indisponível neste servidor.', getSampleStyleSheet()['Normal'])])
                    resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                    return resp

        # Imagens → PDF
        if ext in ('.png', '.jpg', '.jpeg', '.webp', '.gif'):
            try:
                image_to_pdf_reportlab(full_path, preview_path)
                resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                try:
                    resp.headers['Content-Length'] = os.path.getsize(preview_path)
                except Exception:
                    pass
                return resp
            except Exception:
                # Fallback: sempre retorna PDF (erro/minimal) em vez do arquivo original
                try:
                    make_error_pdf(preview_path, 'Conversão não disponível', f'Falha ao converter imagem {stored_name} para PDF.')
                    resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                    try:
                        resp.headers['Content-Length'] = os.path.getsize(preview_path)
                    except Exception:
                        pass
                    return resp
                except Exception:
                    from reportlab.platypus import SimpleDocTemplate, Paragraph
                    from reportlab.lib.pagesizes import A4
                    from reportlab.lib.styles import getSampleStyleSheet
                    doc = SimpleDocTemplate(preview_path, pagesize=A4)
                    doc.build([Paragraph('Conversão de imagem indisponível neste servidor.', getSampleStyleSheet()['Normal'])])
                    resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                    return resp

        # Texto → PDF
        if ext == '.txt':
            try:
                text_to_pdf_reportlab(full_path, preview_path)
                resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                try:
                    resp.headers['Content-Length'] = os.path.getsize(preview_path)
                except Exception:
                    pass
                return resp
            except Exception:
                # Fallback: sempre retorna PDF (erro/minimal)
                try:
                    make_error_pdf(preview_path, 'Conversão não disponível', f'Falha ao converter texto {stored_name} para PDF.')
                    resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                    try:
                        resp.headers['Content-Length'] = os.path.getsize(preview_path)
                    except Exception:
                        pass
                    return resp
                except Exception:
                    from reportlab.platypus import SimpleDocTemplate, Paragraph
                    from reportlab.lib.pagesizes import A4
                    from reportlab.lib.styles import getSampleStyleSheet
                    doc = SimpleDocTemplate(preview_path, pagesize=A4)
                    doc.build([Paragraph('Conversão de texto indisponível neste servidor.', getSampleStyleSheet()['Normal'])])
                    resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                    return resp

        # CSV → PDF
        if ext == '.csv':
            try:
                csv_to_pdf_reportlab(full_path, preview_path)
                resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                try:
                    resp.headers['Content-Length'] = os.path.getsize(preview_path)
                except Exception:
                    pass
                return resp
            except Exception:
                # Fallback: sempre retorna PDF (erro/minimal)
                try:
                    make_error_pdf(preview_path, 'Conversão não disponível', f'Falha ao converter CSV {stored_name} para PDF.')
                    resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                    try:
                        resp.headers['Content-Length'] = os.path.getsize(preview_path)
                    except Exception:
                        pass
                    return resp
                except Exception:
                    from reportlab.platypus import SimpleDocTemplate, Paragraph
                    from reportlab.lib.pagesizes import A4
                    from reportlab.lib.styles import getSampleStyleSheet
                    doc = SimpleDocTemplate(preview_path, pagesize=A4)
                    doc.build([Paragraph('Conversão de CSV indisponível neste servidor.', getSampleStyleSheet()['Normal'])])
                    resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
                    return resp

        # Demais formatos: gera PDF de erro em vez de enviar original
        try:
            make_error_pdf(preview_path, 'Formato não suportado', f'O formato {ext} não é convertido automaticamente. PDF simplificado gerado.')
            resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
            try:
                resp.headers['Content-Length'] = os.path.getsize(preview_path)
            except Exception:
                pass
            return resp
        except Exception:
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            doc = SimpleDocTemplate(preview_path, pagesize=A4)
            doc.build([Paragraph('Formato não suportado para conversão.', getSampleStyleSheet()['Normal'])])
            resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
            return resp
    except Exception as e:
        # Nunca redireciona para HTML: retorna sempre um PDF de erro
        try:
            # Usa o mesmo caminho de preview para gerar o PDF
            preview_dir = ensure_previews_dir()
            preview_name = f'preview_pub_{id_publicacao}.pdf'
            preview_path = os.path.join(preview_dir, preview_name)
            # Nome de download baseado no título
            try:
                safe_title = secure_filename((row.get('titulo') if isinstance(row, dict) else titulo) or 'publicacao')
            except Exception:
                safe_title = 'publicacao'
            download_name = f"{safe_title}.pdf"
            make_error_pdf(preview_path, 'Erro ao preparar download', f'Ocorreu um erro ao gerar o PDF: {e}')
            resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name=download_name)
            try:
                resp.headers['Content-Length'] = os.path.getsize(preview_path)
            except Exception:
                pass
            return resp
        except Exception:
            # Último recurso: PDF mínimo
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            preview_dir = ensure_previews_dir()
            preview_name = f'preview_pub_{id_publicacao}.pdf'
            preview_path = os.path.join(preview_dir, preview_name)
            doc = SimpleDocTemplate(preview_path, pagesize=A4)
            doc.build([Paragraph('Falha ao preparar download em PDF.', getSampleStyleSheet()['Normal'])])
            resp = send_from_directory(preview_dir, preview_name, mimetype='application/pdf', as_attachment=True, download_name='publicacao.pdf')
            return resp

# Rota para reanexar conteúdo de publicação
@app.route('/reupload_publicacao/<int:id_publicacao>', methods=['POST'])
@login_required
@roles_required(['Administrador','Docente'])
def reupload_publicacao(id_publicacao):
    try:
        file = request.files.get('conteudo')
        if not (file and file.filename):
            return jsonify({'ok': False, 'error': 'Arquivo não enviado'}), 400
        ext = os.path.splitext(file.filename)[1].lower()
        allow = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv', '.txt', '.png', '.jpg', '.jpeg', '.webp'}
        if ext not in allow:
            return jsonify({'ok': False, 'error': 'Tipo de arquivo não permitido'}), 400
        filename = secure_filename(file.filename)
        ts = datetime.now().strftime('%Y%m%d%H%M%S')
        new_name = f"{ts}_{filename}"
        upload_dir = app.config['UPLOAD_FOLDER']
        os.makedirs(upload_dir, exist_ok=True)
        full = os.path.join(upload_dir, new_name)
        file.save(full)
        conn = get_db_connection()
        if not conn:
            return jsonify({'ok': False, 'error': 'Falha ao conectar ao banco'}), 500
        # Bloqueia reupload se status não for Pendente
        cur_chk = conn.cursor(row_factory=dict_row)
        cur_chk.execute("SELECT status FROM publicacao WHERE id_publicacao=%s", (id_publicacao,))
        row = cur_chk.fetchone()
        cur_chk.close()
        if not row:
            conn.close()
            return jsonify({'ok': False, 'error': 'Publicação não encontrada'}), 404
        if _norm(str(row.get('status') or '')) != _norm(status_label('Pendente')):
            conn.close()
            return jsonify({'ok': False, 'error': 'Reupload não permitido após avaliação'}), 403
        cur = conn.cursor()
        cur.execute("UPDATE publicacao SET arquivo=%s, nome_arquivo=%s WHERE id_publicacao=%s", (full, new_name, id_publicacao))
        conn.commit()
        cur.close(); conn.close()
        # limpa cache de preview
        preview_dir = ensure_previews_dir()
        preview_name = f'preview_pub_{id_publicacao}.pdf'
        preview_path = os.path.join(preview_dir, preview_name)
        try:
            if os.path.exists(preview_path):
                os.remove(preview_path)
        except Exception:
            pass
        return jsonify({'ok': True, 'filename': new_name})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

# Histórico de avaliações por publicação (JSON)
@app.route('/publicacao/<int:id_publicacao>/avaliacoes', methods=['GET'])
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def publicacao_avaliacoes(id_publicacao):
    conn = get_db_connection()
    if not conn:
        return jsonify({'ok': False, 'error': 'Falha ao conectar ao banco'}), 500
    try:
        cur = conn.cursor(row_factory=dict_row)
        cur.execute(
            """
            SELECT a.id_avaliacao,
                   a.data_avaliacao,
                   a.comentario,
                   u.nome AS avaliador_nome
            FROM avaliacao a
            JOIN usuario u ON u.id_usuario = a.id_avaliador
            WHERE a.id_publicacao = %s
            ORDER BY a.data_avaliacao DESC, a.id_avaliacao DESC
            """,
            (id_publicacao,)
        )
        rows = cur.fetchall()
        cur.close(); conn.close()
        def fmt_row(r):
            dt = r.get('data_avaliacao')
            try:
                dt_str = dt.strftime('%d/%m/%Y %H:%M') if dt else ''
            except Exception:
                dt_str = str(dt) if dt else ''
            return {
                'id_avaliacao': r.get('id_avaliacao'),
                'avaliador': r.get('avaliador_nome') or '',
                'data': dt_str,
                'comentario': r.get('comentario') or ''
            }
        return jsonify({'ok': True, 'avaliacoes': [fmt_row(r) for r in rows]})
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        return jsonify({'ok': False, 'error': str(e)}), 500

# Rota para a página de avaliação
@app.route('/avaliacao')
@login_required
@roles_required(['Administrador','Docente'])
def avaliacao():
    publicacoes = []
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor(row_factory=dict_row)
            # Busca pendentes e denunciados para reavaliação; se Docente, somente as do orientador atual
            is_admin = (session.get('user_tipo') == 'Administrador')
            uid = session.get('user_id')
            pend_label = status_label('Pendente')
            denunc_label = status_label('Denunciado')
            if is_admin:
                cur.execute(
                    """
                    SELECT p.id_publicacao, p.titulo, p.tipo, p.status, p.data_publicacao,
                           u.nome as autor_nome, c.nome_curso, p.id_orientador
                    FROM publicacao p
                    JOIN usuario u ON p.id_autor = u.id_usuario
                    LEFT JOIN curso c ON p.id_curso = c.id_curso
                    WHERE (
                        p.status = %s
                    ) OR (
                        p.status = %s
                        AND NOT EXISTS (
                            SELECT 1 FROM avaliacao a WHERE a.id_publicacao = p.id_publicacao
                        )
                    )
                    ORDER BY p.data_publicacao DESC
                    """,
                    (pend_label, denunc_label)
                )
            else:
                cur.execute(
                    """
                    SELECT p.id_publicacao, p.titulo, p.tipo, p.status, p.data_publicacao,
                           u.nome as autor_nome, c.nome_curso, p.id_orientador
                    FROM publicacao p
                    JOIN usuario u ON p.id_autor = u.id_usuario
                    LEFT JOIN curso c ON p.id_curso = c.id_curso
                    WHERE p.id_orientador = %s AND (
                        p.status = %s
                        OR (
                            p.status = %s
                            AND NOT EXISTS (
                                SELECT 1 FROM avaliacao a WHERE a.id_publicacao = p.id_publicacao
                            )
                        )
                    )
                    ORDER BY p.data_publicacao DESC
                    """,
                    (uid, pend_label, denunc_label)
                )
            publicacoes = cur.fetchall()
            cur.close()
            conn.close()
        except Exception as e:
            flash(f'Erro ao buscar publicações: {e}', 'error')

    return render_template('avaliacao.html', publicacoes=publicacoes)

# Endpoint para decisão de avaliação (Deferir/Indeferir)
@app.route('/avaliacao/decidir', methods=['POST'])
@login_required
@roles_required(['Administrador','Docente'])
def avaliacao_decidir():
    try:
        id_pub = request.form.get('id_publicacao') or (request.json and request.json.get('id_publicacao'))
        try:
            id_pub_int = int(id_pub)
        except Exception:
            id_pub_int = None
        acao = (request.form.get('acao') or (request.json and request.json.get('acao')) or '').strip().lower()
        comentario = (request.form.get('comentario') or (request.json and request.json.get('comentario')) or '').strip()
        if not id_pub_int or acao not in {'deferir','indeferir'}:
            return jsonify({'ok': False, 'error': 'Dados inválidos'}), 400
        conn = get_db_connection()
        if not conn:
            return jsonify({'ok': False, 'error': 'Falha ao conectar ao banco'}), 500
        cur = conn.cursor(row_factory=dict_row)
        cur.execute("SELECT id_orientador, status FROM publicacao WHERE id_publicacao=%s", (id_pub_int,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Publicação não encontrada'}), 404
        # Permitir decisão quando status for 'Pendente' ou 'Denunciado'
        current_status_norm = _norm(str(row.get('status') or ''))
        pend_norm = _norm(status_label('Pendente'))
        denunc_norm = _norm(status_label('Denunciado'))
        if current_status_norm not in {pend_norm, denunc_norm}:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Publicação já avaliada'}), 400
        is_admin = (session.get('user_tipo') == 'Administrador')
        uid = session.get('user_id')
        if not is_admin and int(row.get('id_orientador') or 0) != int(uid or 0):
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Não autorizado para avaliar esta publicação'}), 403
        # Se estiver denunciado e o Docente indeferir, mantém o status como 'Denunciado'
        if acao == 'deferir':
            novo_status = status_label('Publicado')
        else:
            novo_status = status_label('Denunciado') if current_status_norm == denunc_norm else status_label('Reprovado')
        cur2 = conn.cursor()
        # Atualiza status
        cur2.execute("UPDATE publicacao SET status=%s WHERE id_publicacao=%s", (novo_status, id_pub_int))
        # Garante a tabela e insere a avaliação (sem suprimir erro)
        try:
            ensure_avaliacao_table()
        except Exception:
            # Mesmo que a garantia falhe, tentamos inserir e deixamos o erro visível
            pass
        cur2.execute(
            "INSERT INTO avaliacao (id_publicacao, id_avaliador, nota, comentario, data_avaliacao) VALUES (%s,%s,%s,%s,NOW())",
            (id_pub_int, uid, None, comentario or None)
        )
        conn.commit()
        cur2.close(); cur.close(); conn.close()
        nlab = _norm(novo_status)
        if 'public' in nlab:
            ui_status = 'Publicada'
        elif 'denunc' in nlab:
            ui_status = 'Denunciado'
        else:
            ui_status = 'Indeferida'
        return jsonify({'ok': True, 'status': novo_status, 'status_label': ui_status})
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        return jsonify({'ok': False, 'error': str(e)}), 500

# Rota para a página de vinculação de curso
@app.route('/vinculacao_curso', methods=['GET', 'POST'])
@login_required
@roles_required(['Administrador'])
def vinculacao_curso():
    if request.method == 'POST':
        key = f"{request.remote_addr}:vinculacao_curso"
        if not check_rate_limit(key, limit=30, window=60):
            flash('Muitas tentativas. Tente novamente em instantes.', 'error')
            audit_log('rate_limit', {'route': 'vinculacao_curso'})
            return redirect(url_for('vinculacao_curso'))
        usuario_id = request.form.get('usuario')
        curso_id = request.form.get('curso')
        tipo_usuario = request.form.get('tipo_usuario')

        if not usuario_id or not curso_id:
            flash('Selecione curso e usuário para vincular.', 'error')
            return redirect(url_for('vinculacao_curso'))

        conn = get_db_connection()
        if conn:
            try:
                cur = conn.cursor()
                # Garantir vínculo único: remove vínculos anteriores do usuário
                cur.execute("DELETE FROM usuario_curso WHERE id_usuario = %s", (usuario_id,))
                # Insere novo vínculo
                cur.execute(
                    "INSERT INTO usuario_curso (id_usuario, id_curso) VALUES (%s, %s)",
                    (usuario_id, curso_id)
                )
                # Atualiza vínculo direto na tabela de usuário
                cur.execute("UPDATE usuario SET id_curso_usuario = %s WHERE id_usuario = %s", (curso_id, usuario_id))
                conn.commit()
                flash('Vínculo gravado com sucesso.', 'success')
                audit_log('vinculacao_ok', {'usuario_id': usuario_id, 'curso_id': curso_id, 'tipo': tipo_usuario})
                cur.close()
                conn.close()
            except Exception as e:
                flash(f'Erro ao vincular: {e}', 'error')
                audit_log('vinculacao_error', {'error': str(e)})
        return redirect(url_for('vinculacao_curso'))
    
    # Buscar usuários e cursos para o formulário e listar vinculações
    usuarios = []
    cursos = []
    vinculos = []
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor(row_factory=dict_row)
            cur.execute("SELECT * FROM usuario ORDER BY nome")
            usuarios = cur.fetchall()
            
            cur.execute("SELECT * FROM curso ORDER BY nome_curso")
            cursos = cur.fetchall()

            cur.execute(
                """
                SELECT u.id_usuario, u.nome AS usuario, u.tipo AS tipo_usuario,
                       c.id_curso, c.nome_curso AS curso
                FROM usuario u
                JOIN usuario_curso uc ON uc.id_usuario = u.id_usuario
                JOIN curso c ON c.id_curso = uc.id_curso
                ORDER BY u.nome ASC
                LIMIT 100
                """
            )
            vinculos = cur.fetchall()
            
            cur.close()
            conn.close()
        except Exception as e:
            flash(f'Erro ao buscar dados: {e}', 'error')
    
    return render_template('vinculacao_curso.html', usuarios=usuarios, cursos=cursos, vinculos=vinculos)

# API para listar usuários por tipo
@app.route('/api/usuarios_por_tipo/<tipo>', methods=['GET'])
@login_required
@roles_required(['Administrador'])
def api_usuarios_por_tipo(tipo):
    data = []
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor(row_factory=dict_row)
            cur.execute("SELECT id_usuario, nome FROM usuario WHERE tipo = %s AND ativo = TRUE ORDER BY nome", (tipo,))
            data = cur.fetchall()
            cur.close(); conn.close()
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify(data)

# Remover vínculo de curso do usuário
@app.route('/vinculacao_curso/remover', methods=['POST'])
@login_required
@roles_required(['Administrador'])
def remover_vinculo_curso():
    usuario_id = request.form.get('usuario_id')
    if not usuario_id:
        flash('Usuário inválido para remover vínculo.', 'error')
        return redirect(url_for('vinculacao_curso'))
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor()
            cur.execute("DELETE FROM usuario_curso WHERE id_usuario = %s", (usuario_id,))
            cur.execute("UPDATE usuario SET id_curso_usuario = NULL WHERE id_usuario = %s", (usuario_id,))
            conn.commit()
            cur.close(); conn.close()
            flash('Vínculo removido.', 'success')
            audit_log('vinculo_removido', {'usuario_id': usuario_id})
        except Exception as e:
            flash(f'Erro ao remover vínculo: {e}', 'error')
            audit_log('vinculo_remover_error', {'error': str(e)})
    return redirect(url_for('vinculacao_curso'))

# Rota para a página de relatório
@app.route('/relatorio')
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def relatorio():
    autores = []
    cursos = []
    tipos = []
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor(row_factory=dict_row)
            # Orientadores/Professores: listar todos os usuários com perfil Professor/Docente
            cur.execute("""
                SELECT u.id_usuario, u.nome
                FROM usuario u
                WHERE u.tipo::text IN ('Professor','Docente')
                ORDER BY u.nome
            """)
            autores = cur.fetchall()
            # Cursos: todos os cursos
            cur.execute("SELECT id_curso, nome_curso FROM curso ORDER BY nome_curso")
            cursos = cur.fetchall()
            # Tipos de publicação aceitos
            cur.execute("SELECT nome_tipo FROM tipos_de_publicacao ORDER BY nome_tipo")
            tipos = cur.fetchall()
            cur.close()
            conn.close()
        except Exception as e:
            try:
                conn.close()
            except Exception:
                pass
            flash(f'Erro ao carregar filtros de relatório: {e}', 'error')
    return render_template('relatorio.html', autores=autores, cursos=cursos, tipos=tipos)

# Exportação de Relatório em Excel
@app.route('/relatorio/exportar', methods=['GET'])
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def exportar_relatorio():
    # Coleta de filtros
    autor = (request.args.get('autor') or '').strip()
    orientador = (request.args.get('orientador') or '').strip()
    curso = (request.args.get('curso') or '').strip()
    tipo = (request.args.get('tipo') or '').strip()
    status = (request.args.get('status') or '').strip()
    data_inicial = (request.args.get('data_inicial') or '').strip()
    data_final = (request.args.get('data_final') or '').strip()

    # Construção dinâmica do WHERE
    where = ["1=1"]
    params = []

    if autor:
        where.append("u.nome ILIKE %s")
        params.append(f"%{autor}%")
    if orientador:
        where.append("p.id_orientador = %s")
        params.append(orientador)
    if curso:
        where.append("p.id_curso = %s")
        params.append(curso)
    if tipo:
        where.append("p.tipo = %s")
        params.append(tipo)
    # Status
    if status:
        valid_status = {'Pendente','Publicado','Denunciado'}
        if status in valid_status:
            where.append("p.status = %s")
            params.append(status)
    # Datas (YYYY-MM-DD)
    try:
        from datetime import datetime
        if data_inicial:
            di = None
            for fmt in ('%d/%m/%Y', '%Y-%m-%d'):
                try:
                    di = datetime.strptime(data_inicial, fmt).date()
                    break
                except Exception:
                    pass
            if di is not None:
                where.append("p.data_publicacao >= %s")
                params.append(di)
        if data_final:
            df = None
            for fmt in ('%d/%m/%Y', '%Y-%m-%d'):
                try:
                    df = datetime.strptime(data_final, fmt).date()
                    break
                except Exception:
                    pass
            if df is not None:
                where.append("p.data_publicacao <= %s")
                params.append(df)
    except Exception:
        pass

    sql = f"""
        SELECT 
          p.id_publicacao,
          p.titulo,
          p.tipo,
          p.data_publicacao,
          p.status,
          COALESCE(u_orient.nome, u_autor.nome, '') AS autor,
          COALESCE(c.nome_curso, '') AS curso,
          COALESCE(p.assuntos_relacionados, '') AS assuntos
        FROM publicacao p
        LEFT JOIN usuario u_autor ON u_autor.id_usuario = p.id_autor
        LEFT JOIN usuario u_orient ON u_orient.id_usuario = p.id_orientador
        LEFT JOIN curso c ON c.id_curso = p.id_curso
        WHERE {' AND '.join(where)}
        ORDER BY p.data_publicacao DESC, p.id_publicacao DESC
    """

    try:
        conn = get_db_connection()
        if not conn:
            flash('Falha ao conectar para exportação.', 'error')
            return redirect(url_for('relatorio'))
        cur = conn.cursor(row_factory=dict_row)
        cur.execute(sql, params)
        rows = cur.fetchall() or []
        cur.close(); conn.close()

        # Limite de taxa simples por usuário (10/min)
        key = f"relatorio_export::{session.get('user_id') or 'anon'}"
        if not check_rate_limit(key, limit=10, window=60):
            return make_response('Muitas exportações. Tente novamente em instantes.', 429)

        # Colunas configuráveis
        import re, io
        fmt = (request.args.get('format') or 'xlsx').lower()
        cols_param = (request.args.get('cols') or '').strip()
        valid_cols = ['id_publicacao','titulo','tipo','autor','curso','data_publicacao','status','assuntos']
        col_map = {
            'id_publicacao':'ID',
            'titulo':'Título',
            'tipo':'Tipo',
            'autor':'Orientador/ Professor',
            'curso':'Curso',
            'data_publicacao':'Data Publicação',
            'status':'Status',
            'assuntos':'Assuntos'
        }
        selected_cols = [c for c in re.split(r'[\s,;]+', cols_param) if c in valid_cols]
        if not selected_cols:
            selected_cols = valid_cols[:]

        from datetime import datetime as _dt
        fname_base = f"relatorio_inprolib_{_dt.now().strftime('%Y%m%d_%H%M%S')}"

        def val_for(col_key, r):
            v = r.get(col_key)
            if col_key == 'data_publicacao' and v:
                try:
                    # psycopg returns date/datetime; para CSV/PDF usar dd/mm/aaaa
                    return v.strftime('%d/%m/%Y')
                except Exception:
                    return str(v)
            return v if (v is not None) else ''

        if fmt == 'csv':
            import csv
            headers = [col_map[c] for c in selected_cols]
            text_buf = io.StringIO()
            writer = csv.writer(text_buf, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL)
            writer.writerow(headers)
            for r in rows:
                writer.writerow([val_for(c, r) for c in selected_cols])
            csv_bytes = text_buf.getvalue().encode('utf-8-sig')  # BOM para Excel abrir corretamente
            buf = io.BytesIO(csv_bytes)
            buf.seek(0)
            resp = send_file(
                buf,
                as_attachment=True,
                download_name=f"{fname_base}.csv",
                mimetype='text/csv',
                max_age=0
            )
        elif fmt == 'pdf':
            try:
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib.pagesizes import landscape, A4
                from reportlab.lib import colors
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                headers = [col_map[c] for c in selected_cols]
                buf = io.BytesIO()
                doc = SimpleDocTemplate(buf, pagesize=landscape(A4), rightMargin=24, leftMargin=24, topMargin=24, bottomMargin=24)
                styles = getSampleStyleSheet()
                title_style = ParagraphStyle('ReportTitle', parent=styles['Title'], fontSize=16, textColor=colors.HexColor('#1F4E79'), alignment=1, spaceAfter=12)
                title = Paragraph('Relatório de publicações', title_style)
                data = [headers]
                for r in rows:
                    data.append([val_for(c, r) for c in selected_cols])
                table = Table(data, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#f0f0f0')),
                    ('TEXTCOLOR',(0,0),(-1,0),colors.black),
                    ('GRID',(0,0),(-1,-1),0.25,colors.HexColor('#cccccc')),
                    ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
                    ('FONTSIZE',(0,0),(-1,-1),9),
                    ('ALIGN',(0,0),(-1,-1),'LEFT'),
                    ('VALIGN',(0,0),(-1,-1),'MIDDLE')
                ]))
                doc.build([title, Spacer(1, 10), table])
                buf.seek(0)
                resp = send_file(
                    buf,
                    as_attachment=True,
                    download_name=f"{fname_base}.pdf",
                    mimetype='application/pdf',
                    max_age=0
                )
            except Exception as _e:
                # Fallback simples para CSV se PDF falhar
                import csv
                headers = [col_map[c] for c in selected_cols]
                text_buf = io.StringIO()
                writer = csv.writer(text_buf, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL)
                writer.writerow(headers)
                for r in rows:
                    writer.writerow([val_for(c, r) for c in selected_cols])
                csv_bytes = text_buf.getvalue().encode('utf-8-sig')
                buf = io.BytesIO(csv_bytes)
                buf.seek(0)
                resp = send_file(
                    buf,
                    as_attachment=True,
                    download_name=f"{fname_base}.csv",
                    mimetype='text/csv',
                    max_age=0
                )
        else:
            # Excel sem a coluna ID e com formatação moderna
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            wb = Workbook()
            ws = wb.active
            ws.title = 'Relatório'
            excel_cols = [c for c in selected_cols if c != 'id_publicacao']
            if not excel_cols:
                excel_cols = [c for c in valid_cols if c != 'id_publicacao']
            excel_headers = [col_map[c] for c in excel_cols]
            # Título do relatório
            title_text = 'Relatório de publicações'
            ws.append([title_text])
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(excel_headers))
            ws['A1'].font = Font(size=16, bold=True, color='1F4E79')
            ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
            ws.row_dimensions[1].height = 28
            hdr_row = 2
            # Cabeçalhos
            ws.append(excel_headers)
            for r in rows:
                row_vals = []
                for c in excel_cols:
                    v = r.get(c)
                    if c == 'data_publicacao' and v:
                        row_vals.append(v)  # manter como date/datetime para formatar
                    else:
                        row_vals.append(v if (v is not None) else '')
                ws.append(row_vals)
            # Estilo moderno
            thin = Side(style='thin', color='D0D7DE')
            hdr_fill = PatternFill('solid', fgColor='F1F5F9')
            for col_idx in range(1, len(excel_headers)+1):
                cell = ws.cell(row=hdr_row, column=col_idx)
                cell.font = Font(bold=True)
                cell.fill = hdr_fill
                cell.alignment = Alignment(vertical='center')
                cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
            # Zebrado
            for i in range(0, len(rows)):
                if i % 2 == 0:
                    row_num = i + hdr_row + 1
                    for col_idx in range(1, len(excel_headers)+1):
                        ws.cell(row=row_num, column=col_idx).fill = PatternFill('solid', fgColor='F9FAFB')
            # Wrap em Título/Assuntos
            wrap_cols = []
            for idx, key in enumerate(excel_cols, start=1):
                if key in {'titulo','assuntos'}:
                    wrap_cols.append(idx)
            for row_num in range(hdr_row+1, ws.max_row+1):
                for idx in wrap_cols:
                    ws.cell(row=row_num, column=idx).alignment = Alignment(wrap_text=True, vertical='top')
            # Larguras de coluna
            width_map = {
                'titulo': 50,
                'tipo': 18,
                'autor': 24,
                'curso': 24,
                'data_publicacao': 14,
                'status': 16,
                'assuntos': 36
            }
            for idx, key in enumerate(excel_cols, start=1):
                ws.column_dimensions[get_column_letter(idx)].width = width_map.get(key, 22)
            # Congelar cabeçalho e filtro
            ws.freeze_panes = f'A{hdr_row+1}'
            ws.auto_filter.ref = f"A{hdr_row}:{get_column_letter(len(excel_headers))}{hdr_row}"
            # Formato de data
            if 'data_publicacao' in excel_cols:
                d_idx = excel_cols.index('data_publicacao') + 1
                for row_num in range(hdr_row+1, ws.max_row+1):
                    ws.cell(row=row_num, column=d_idx).number_format = 'DD/MM/YYYY'
            buf = io.BytesIO()
            wb.save(buf)
            buf.seek(0)
            resp = send_file(
                buf,
                as_attachment=True,
                download_name=f"{fname_base}.xlsx",
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                max_age=0
            )

        # Define Content-Length explícito para permitir barra de progresso no frontend
        try:
            resp.headers['Content-Length'] = buf.getbuffer().nbytes
        except Exception:
            pass
        resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        resp.headers['Pragma'] = 'no-cache'
        resp.headers['Expires'] = '0'
        # Auditoria leve
        try:
            audit_log('relatorio_export', {
                'format': fmt,
                'cols': selected_cols,
                'rows': len(rows),
                'filters': {
                    'autor': autor,
                    'orientador': orientador,
                    'curso': curso,
                    'tipo': tipo,
                    'status': status,
                    'data_inicial': data_inicial,
                    'data_final': data_final
                }
            })
        except Exception:
            pass
        return resp
    except Exception as e:
        flash(f'Erro ao gerar Excel: {e}', 'error')
        try:
            conn.close()
        except Exception:
            pass
        return redirect(url_for('relatorio'))

# Visualização de Relatório (JSON)
@app.route('/relatorio/preview', methods=['GET'])
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def preview_relatorio():
    autor = (request.args.get('autor') or '').strip()
    orientador = (request.args.get('orientador') or '').strip()
    curso = (request.args.get('curso') or '').strip()
    tipo = (request.args.get('tipo') or '').strip()
    status = (request.args.get('status') or '').strip()
    data_inicial = (request.args.get('data_inicial') or '').strip()
    data_final = (request.args.get('data_final') or '').strip()

    where = ["1=1"]
    params = []
    if autor:
        where.append("u.nome ILIKE %s")
        params.append(f"%{autor}%")
    if orientador:
        where.append("p.id_orientador = %s")
        params.append(orientador)
    if curso:
        where.append("p.id_curso = %s")
        params.append(curso)
    if tipo:
        where.append("p.tipo = %s")
        params.append(tipo)
    if status:
        valid_status = {'Pendente','Publicado','Denunciado'}
        if status in valid_status:
            where.append("p.status = %s")
            params.append(status)
    try:
        from datetime import datetime
        if data_inicial:
            di = None
            for fmt in ('%d/%m/%Y', '%Y-%m-%d'):
                try:
                    di = datetime.strptime(data_inicial, fmt).date()
                    break
                except Exception:
                    pass
            if di is not None:
                where.append("p.data_publicacao >= %s")
                params.append(di)
        if data_final:
            df = None
            for fmt in ('%d/%m/%Y', '%Y-%m-%d'):
                try:
                    df = datetime.strptime(data_final, fmt).date()
                    break
                except Exception:
                    pass
            if df is not None:
                where.append("p.data_publicacao <= %s")
                params.append(df)
    except Exception:
        pass

    where_clause = " AND ".join(where)
    sql = f"""
        SELECT 
          p.id_publicacao, p.titulo, p.tipo, p.status, p.assuntos_relacionados as assuntos,
          COALESCE(u_orient.nome, u_autor.nome) as autor, c.nome_curso as curso, p.data_publicacao
        FROM publicacao p
        LEFT JOIN usuario u_autor ON u_autor.id_usuario = p.id_autor
        LEFT JOIN usuario u_orient ON u_orient.id_usuario = p.id_orientador
        LEFT JOIN curso c ON c.id_curso = p.id_curso
        WHERE {where_clause}
        ORDER BY p.id_publicacao DESC
    """

    try:
        conn = get_db_connection()
        if not conn:
            return make_response(jsonify({'error': 'Falha ao conectar ao banco.'}), 500)
        cur = conn.cursor(row_factory=dict_row)
        cur.execute(sql, params)
        rows = cur.fetchall() or []
        cur.close(); conn.close()
        # Normaliza datas para string
        out = []
        for r in rows:
            out.append({
                'id_publicacao': r.get('id_publicacao'),
                'titulo': r.get('titulo'),
                'tipo': r.get('tipo'),
                'autor': r.get('autor'),
                'curso': r.get('curso'),
                'data_publicacao': r.get('data_publicacao').strftime('%d/%m/%Y') if r.get('data_publicacao') else '',
                'status': r.get('status'),
                'assuntos': r.get('assuntos')
            })
        return jsonify({'rows': out})
    except Exception as e:
        return make_response(jsonify({'error': str(e)}), 500)

# Rota para a página de suporte
@app.route('/suporte', methods=['GET','POST'])
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def suporte():
    if request.method == 'POST':
        key = f"{request.remote_addr}:suporte"
        if not check_rate_limit(key, limit=10, window=60):
            flash('Muitas tentativas. Tente novamente em instantes.', 'error')
            audit_log('rate_limit', {'route': 'suporte'})
            return redirect(url_for('suporte'))
        mensagem = (request.form.get('mensagem') or '').strip()
        arquivo = request.files.get('imagem')
        attach_tuple = None
        try:
            if arquivo and arquivo.filename:
                filename = secure_filename(arquivo.filename)
                data_bytes = arquivo.read()
                mimetype = arquivo.mimetype or (mimetypes.guess_type(filename)[0] or 'application/octet-stream')
                attach_tuple = (filename, data_bytes, mimetype)
        except Exception:
            attach_tuple = None
        user_name = (session.get('user_name') or '').strip()
        user_id = session.get('user_id')
        role = (session.get('role') or session.get('user_tipo') or '').strip()
        # Validação: mensagem obrigatória (front já impõe required, mas reforçamos no backend)
        if not mensagem:
            flash('Por favor, descreva sua mensagem para o suporte.', 'error')
            return redirect(url_for('suporte'))

        # Segurança: garantir limite máximo de 3200 caracteres
        mensagem_limpa = mensagem.strip()[:3200]

        # Tentar obter e-mail do usuário para Reply-To
        reply_email = None
        try:
            conn = get_db_connection()
            if conn and user_id:
                cur = conn.cursor(row_factory=dict_row)
                cur.execute('SELECT email FROM usuario WHERE id_usuario = %s', (user_id,))
                row = cur.fetchone()
                if row and row.get('email'):
                    reply_email = (row.get('email') or '').strip()
                cur.close(); conn.close()
        except Exception:
            try:
                conn and conn.close()
            except Exception:
                pass

        body_text = (
            'Novo contato de suporte no INPROLIB:\n\n'
            f'Usuário: {user_name or "Desconhecido"}\n'
            f'Perfil: {role or "-"}\n'
            f'ID do usuário: {user_id or "-"}\n\n'
            'Mensagem do usuário:\n'
            f'{mensagem_limpa}\n'
        )
        ok = send_support_email(body_text, attach_tuple, reply_email)
        if ok:
            flash('Mensagem de suporte enviada com sucesso!', 'success')
            audit_log('suporte_email_ok', {'user_id': user_id})
        else:
            flash('Falha ao enviar o e-mail de suporte.', 'error')
            audit_log('suporte_email_error', {'user_id': user_id})
        return redirect(url_for('suporte'))
    return render_template('suporte.html')

# Denúncia de publicação: volta para avaliação do orientador com status 'Denunciado'
@app.route('/publicacao/denuncia', methods=['POST'])
@login_required
@roles_required(['Administrador','Aluno'])
def publicacao_denuncia():
    try:
        key = f"{request.remote_addr}:denuncia"
        if not check_rate_limit(key, limit=10, window=60):
            audit_log('rate_limit', {'route': 'publicacao_denuncia'})
            return jsonify({'ok': False, 'error': 'Muitas tentativas. Tente novamente em instantes.'}), 429

        id_pub = request.form.get('id_publicacao') or (request.json and request.json.get('id_publicacao'))
        try:
            id_pub_int = int(id_pub)
        except Exception:
            id_pub_int = None
        mensagem_raw = (request.form.get('mensagem') or '').strip()
        # Mensagem simples enviada separadamente pelo front para validação de 800 caracteres
        mensagem_plain = (request.form.get('mensagem_plain') or '').strip()
        if not mensagem_plain:
            # tenta extrair a parte após 'Descrição do usuário:' do texto completo
            try:
                idx = mensagem_raw.lower().rfind('descrição do usuário:')
                if idx != -1:
                    mensagem_plain = mensagem_raw[idx+len('descrição do usuário:'):].strip()
                else:
                    mensagem_plain = mensagem_raw
            except Exception:
                mensagem_plain = mensagem_raw
        # aplica limite de 800 caracteres para a descrição
        if len(mensagem_plain) > 800:
            return jsonify({'ok': False, 'error': 'Descrição da denúncia deve ter no máximo 800 caracteres.'}), 400
        # mensagem final a ser enviada por e-mail (com contexto enxuto)
        mensagem = mensagem_plain
        arquivo = request.files.get('imagem')
        attach_tuple = None
        try:
            if arquivo and arquivo.filename:
                filename = secure_filename(arquivo.filename)
                data_bytes = arquivo.read()
                mimetype = arquivo.mimetype or (mimetypes.guess_type(filename)[0] or 'application/octet-stream')
                attach_tuple = (filename, data_bytes, mimetype)
        except Exception:
            attach_tuple = None

        if not id_pub_int:
            return jsonify({'ok': False, 'error': 'Publicação inválida.'}), 400

        # exige imagem de evidência
        if not attach_tuple:
            return jsonify({'ok': False, 'error': 'Anexe uma imagem do erro para avaliação.'}), 400

        # Atualiza status para Denunciado para aparecer na tela de avaliação
        conn = get_db_connection()
        if not conn:
            return jsonify({'ok': False, 'error': 'Falha ao conectar ao banco.'}), 500
        try:
            denunc = status_label('Denunciado')
            with conn.cursor(row_factory=dict_row) as cur:
                cur.execute("SELECT id_publicacao, titulo, id_orientador FROM publicacao WHERE id_publicacao=%s", (id_pub_int,))
                row = cur.fetchone()
                if not row:
                    conn.close()
                    return jsonify({'ok': False, 'error': 'Publicação não encontrada.'}), 404
                cur.execute("UPDATE publicacao SET status=%s WHERE id_publicacao=%s", (denunc, id_pub_int))
                conn.commit()
            conn.close()
        except Exception as e:
            try:
                conn and conn.close()
            except Exception:
                pass
            return jsonify({'ok': False, 'error': f'Falha ao atualizar status: {e}'}), 500

        # Envia e-mail ao suporte com contexto (mantém padrão existente de suporte)
        user_name = (session.get('user_name') or '').strip()
        role = (session.get('role') or session.get('user_tipo') or '').strip()
        body_text = (
            'Denúncia de COPYRIGHT no INPROLIB:\n\n'
            f'Usuário: {user_name or "Desconhecido"}\n'
            f'Perfil: {role or "-"}\n'
            f'ID da publicação: {id_pub_int}\n\n'
            'Descrição do usuário:\n'
            f'{mensagem or "(sem descrição)"}\n'
        )
        try:
            send_support_email(body_text, attach_tuple, None, subject='INPROLIB - Denúncia de publicação')
        except Exception:
            pass
        audit_log('denuncia_publicacao', {'id_publicacao': id_pub_int})
        return jsonify({'ok': True, 'redirect': '/publicacao'})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

# Rota para a página de configuração
@app.route('/configuracao')
@login_required
@roles_required(['Administrador'])
def configuracao():
    uid = session.get('user_id')
    user = None
    usuarios = []
    try:
        conn = get_db_connection()
        if conn and uid:
            cur = conn.cursor(row_factory=dict_row)
            # Usuário inicialmente selecionado (admin começa com ele mesmo)
            cur.execute('SELECT id_usuario, nome, cpf, email, telefone, tema_preferido FROM usuario WHERE id_usuario = %s', (uid,))
            user = cur.fetchone()
            # Lista de usuários para seleção/edição
            cur.execute('SELECT id_usuario, nome FROM usuario ORDER BY nome ASC')
            usuarios = cur.fetchall()
            cur.close(); conn.close()
    except Exception:
        try:
            conn and conn.close()
        except Exception:
            pass
    return render_template('configuracao.html', user=user, usuarios=usuarios)

# Atualização de perfil
@app.route('/configuracao/perfil', methods=['POST'])
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def configuracao_perfil():
    uid = session.get('user_id')
    role = (session.get('role') or session.get('user_tipo') or '')
    nome = (request.form.get('nome') or '').strip()
    email = (request.form.get('email') or '').strip().lower()
    # telefone tornou-se opcional no formulário; não atualiza se não vier
    telefone_present = ('telefone' in request.form)
    telefone = ((request.form.get('telefone') or '').strip() if telefone_present else None)
    # Permite que Administrador edite outro usuário selecionado
    target_uid = uid
    if role == 'Administrador':
        id_usuario_raw = request.form.get('id_usuario')
        try:
            if id_usuario_raw:
                target_uid = int(id_usuario_raw)
        except Exception:
            pass
    if not nome or not email:
        msg = 'Nome e e-mail são obrigatórios.'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'ok': False, 'error': msg}), 400
        flash(msg, 'error'); return redirect(url_for('configuracao'))
    if '@' not in email or '.' not in email:
        msg = 'E-mail inválido.'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'ok': False, 'error': msg}), 400
        flash(msg, 'error'); return redirect(url_for('configuracao'))
    try:
        conn = get_db_connection()
        if not conn:
            raise Exception('Falha de conexão')
        cur = conn.cursor()
        cur.execute("SELECT id_usuario FROM usuario WHERE LOWER(email) = %s AND id_usuario <> %s", (email, target_uid))
        dup = cur.fetchone()
        if dup:
            cur.close(); conn.close()
            msg = 'E-mail já cadastrado por outro usuário.'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'ok': False, 'error': msg}), 409
            flash(msg, 'error'); return redirect(url_for('configuracao'))
        if telefone_present:
            cur.execute("UPDATE usuario SET nome = %s, email = %s, telefone = %s WHERE id_usuario = %s", (nome, email, telefone, target_uid))
        else:
            cur.execute("UPDATE usuario SET nome = %s, email = %s WHERE id_usuario = %s", (nome, email, target_uid))
        conn.commit(); cur.close(); conn.close()
        # Só atualiza nome na sessão se o usuário editado for o próprio
        if target_uid == uid:
            session['user_name'] = nome
        msg = 'Perfil atualizado com sucesso.'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'ok': True, 'message': msg})
        flash(msg, 'success'); return redirect(url_for('configuracao'))
    except Exception as e:
        try:
            conn and conn.close()
        except Exception:
            pass
        msg = f'Erro ao atualizar perfil: {e}'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'ok': False, 'error': msg}), 500
        flash(msg, 'error'); return redirect(url_for('configuracao'))

# Detalhes de um usuário (para preencher formulário de edição)
@app.route('/api/usuario/<int:user_id>', methods=['GET'])
@login_required
@roles_required(['Administrador'])
def api_usuario(user_id: int):
    try:
        conn = get_db_connection()
        if not conn:
            raise Exception('Falha de conexão')
        cur = conn.cursor(row_factory=dict_row)
        cur.execute('SELECT id_usuario, nome, email, telefone FROM usuario WHERE id_usuario = %s', (user_id,))
        row = cur.fetchone()
        cur.close(); conn.close()
        if not row:
            return jsonify({'ok': False, 'error': 'Usuário não encontrado.'}), 404
        return jsonify({'ok': True, 'user': row})
    except Exception as e:
        try:
            conn and conn.close()
        except Exception:
            pass
        return jsonify({'ok': False, 'error': str(e)}), 500

# Alteração de senha
@app.route('/configuracao/senha', methods=['POST'])
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def configuracao_senha():
    uid = session.get('user_id')
    atual = (request.form.get('senha_atual') or '').strip()
    nova = (request.form.get('nova_senha') or '').strip()
    conf = (request.form.get('confirmar_senha') or '').strip()
    if not atual or not nova or not conf:
        msg = 'Preencha todos os campos de senha.'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'ok': False, 'error': msg}), 400
        flash(msg, 'error'); return redirect(url_for('configuracao'))
    if nova != conf:
        msg = 'Nova senha e confirmação não coincidem.'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'ok': False, 'error': msg}), 400
        flash(msg, 'error'); return redirect(url_for('configuracao'))
    if not password_policy_ok(nova):
        msg = 'A nova senha deve ter exatamente 8 caracteres, com maiúscula, minúscula, número e símbolo.'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'ok': False, 'error': msg}), 400
        flash(msg, 'error'); return redirect(url_for('configuracao'))
    try:
        conn = get_db_connection()
        if not conn:
            raise Exception('Falha de conexão')
        cur = conn.cursor(row_factory=dict_row)
        cur.execute('SELECT senha FROM usuario WHERE id_usuario = %s', (uid,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            msg = 'Usuário inválido.'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'ok': False, 'error': msg}), 404
            flash(msg, 'error'); return redirect(url_for('configuracao'))
        senha_hash = row['senha'] if isinstance(row['senha'], str) else (row['senha'].decode() if row['senha'] else '')
        if not senha_hash or not check_password_hash(senha_hash, atual):
            cur.close(); conn.close()
            msg = 'Senha atual incorreta.'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'ok': False, 'error': msg}), 400
            flash(msg, 'error'); return redirect(url_for('configuracao'))
        novo_hash = generate_password_hash(nova)
        cur2 = conn.cursor()
        cur2.execute('UPDATE usuario SET senha = %s WHERE id_usuario = %s', (novo_hash, uid))
        conn.commit(); cur2.close(); cur.close(); conn.close()
        msg = 'Senha atualizada com sucesso.'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'ok': True, 'message': msg})
        flash(msg, 'success'); return redirect(url_for('configuracao'))
    except Exception as e:
        try:
            conn and conn.close()
        except Exception:
            pass
        msg = f'Erro ao atualizar senha: {e}'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'ok': False, 'error': msg}), 500
        flash(msg, 'error'); return redirect(url_for('configuracao'))

# Alternância de tema
@app.route('/configuracao/tema', methods=['POST'])
@login_required
@roles_required(['Administrador','Docente','Aluno'])
def configuracao_tema():
    uid = session.get('user_id')
    tema = (request.form.get('tema') or '').strip().lower()
    if tema not in ('claro','escuro'):
        return jsonify({'ok': False, 'error': 'Tema inválido.'}), 400
    try:
        conn = get_db_connection()
        if not conn:
            raise Exception('Falha de conexão')
        cur = conn.cursor()
        cur.execute('UPDATE usuario SET tema_preferido = %s WHERE id_usuario = %s', (tema, uid))
        conn.commit(); cur.close(); conn.close()
        session['user_theme'] = tema
        return jsonify({'ok': True, 'tema': tema})
    except Exception as e:
        try:
            conn and conn.close()
        except Exception:
            pass
        return jsonify({'ok': False, 'error': f'Erro ao salvar tema: {e}'}), 500

# Rota para logout
@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

# API de CEP (proxy ViaCEP)
@app.route('/api/cep/<cep>', methods=['GET'])
def api_cep(cep):
    key = f"{request.remote_addr}:api_cep"
    if not check_rate_limit(key, limit=40, window=60):
        return make_response(jsonify({'erro': 'rate_limit'}), 429)
    digits = re.sub(r'[^0-9]', '', (cep or ''))
    if len(digits) != 8:
        return make_response(jsonify({'erro': 'CEP inválido'}), 400)
    import urllib.request
    import ssl
    url = f"https://viacep.com.br/ws/{digits}/json/"
    try:
        ctx = ssl.create_default_context()
        with urllib.request.urlopen(url, timeout=5, context=ctx) as resp:
            payload = resp.read().decode('utf-8')
        data = json.loads(payload or '{}')
        if data.get('erro'):
            return make_response(jsonify({'erro': 'CEP não encontrado'}), 404)
        out = {
            'cep': data.get('cep', ''),
            'logradouro': data.get('logradouro', ''),
            'complemento': data.get('complemento', ''),
            'bairro': data.get('bairro', ''),
            'localidade': data.get('localidade', ''),
            'uf': data.get('uf', '')
        }
        return jsonify(out)
    except Exception as e:
        print(f'[CEP] Falha ao consultar ViaCEP: {e}')
        return make_response(jsonify({'erro': 'Falha ao consultar CEP'}), 502)

# API para buscar publicações
@app.route('/api/publicacoes', methods=['GET'])
def api_publicacoes():
    query = request.args.get('q', '')
    filtros = request.args.getlist('filtro')
    
    if not filtros:
        filtros = ['autor', 'assunto', 'curso', 'titulo']
    
    conn = get_db_connection()
    resultados = []
    
    if conn:
        try:
            cur = conn.cursor(row_factory=dict_row)
            
            # Construir a consulta SQL com base nos filtros
            sql_parts = []
            params = []
            
            if 'autor' in filtros:
                sql_parts.append("u.nome ILIKE %s")
                params.append(f'%{query}%')
            
            if 'assunto' in filtros:
                sql_parts.append("p.assuntos_relacionados ILIKE %s")
                params.append(f'%{query}%')
            
            if 'curso' in filtros:
                sql_parts.append("c.nome_curso ILIKE %s")
                params.append(f'%{query}%')
            
            if 'titulo' in filtros:
                sql_parts.append("p.titulo ILIKE %s")
                params.append(f'%{query}%')
            
            where_clause = " OR ".join(sql_parts)
            
            # Executar a consulta
            cur.execute(f"""
                SELECT p.*, u.nome as autor_nome, c.nome_curso 
                FROM publicacao p
                JOIN usuario u ON p.id_autor = u.id_usuario
                JOIN curso c ON p.id_curso = c.id_curso
                WHERE p.status = 'Publicado' AND ({where_clause})
                ORDER BY p.data_publicacao DESC
            """, params)
            
            resultados = cur.fetchall()
            
            # Converter para formato JSON
            resultados_json = []
            for r in resultados:
                # Padroniza formato de data para dd/mm/aaaa
                try:
                    dt_str = r['data_publicacao'].strftime('%d/%m/%Y') if r['data_publicacao'] else ''
                except Exception:
                    dt_str = str(r['data_publicacao']) if r.get('data_publicacao') else ''
                resultados_json.append({
                    'id': r['id_publicacao'],
                    'titulo': r['titulo'],
                    'autor': r['autor_nome'],
                    'curso': r['nome_curso'],
                    'data': dt_str,
                    'tipo': r['tipo'],
                    'assuntos': r['assuntos_relacionados']
                })
            
            cur.close()
            conn.close()
            return jsonify(resultados_json)
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    return jsonify([])

def run_validacao():
    # Rotina de validação: cria curso, publica arquivo com CAPTCHA e vincula usuário a curso.
    with app.test_client() as client:
        # 1) Criar curso
        resp1 = client.post('/cadastro_curso', data={
            'nome_curso': 'Curso Validação',
            'descricao': 'Curso de Teste',
            'codigo': 'VAL123',
            'autorizacao': '1234',
            'coordenador': ''
        }, follow_redirects=True)
        print('cadastro_curso status:', resp1.status_code)

        # Obter id do curso recém-criado (ou último com mesmo código)
        curso_id = None
        conn = get_db_connection()
        if conn:
            try:
                cur = conn.cursor()
                cur.execute("""
                    SELECT id_curso FROM curso
                    WHERE codigo_curso = %s
                    ORDER BY id_curso DESC
                    LIMIT 1
                """, ('VAL123',))
                row = cur.fetchone()
                if row:
                    curso_id = str(row[0])
                cur.close()
                conn.close()
            except Exception as e:
                try:
                    conn.close()
                except Exception:
                    pass
                print('Erro ao obter curso_id:', e)

        # 2) Publicação com CAPTCHA e upload
        _ = client.get('/publicacao')
        with client.session_transaction() as sess:
            captcha_answer = sess.get('captcha_answer')
        data = {
            'titulo_conteudo': 'Teste Publicacao',
            'tipo_publicacao': 'TCC',
            'curso': curso_id or '',
            'captcha': captcha_answer or ''
        }
        file_tuple = (io.BytesIO(b'Arquivo de teste da publicacao'), 'test_upload.txt')
        resp2 = client.post(
            '/publicacao',
            data={**data, 'conteudo': file_tuple},
            content_type='multipart/form-data',
            follow_redirects=True
        )
        print('publicacao status:', resp2.status_code)

        # 3) Vincular usuário a curso (pega primeiro usuário existente)
        usuario_id = None
        conn2 = get_db_connection()
        if conn2:
            try:
                cur2 = conn2.cursor()
                cur2.execute("SELECT id_usuario FROM usuario ORDER BY id_usuario LIMIT 1")
                row2 = cur2.fetchone()
                if row2:
                    usuario_id = str(row2[0])
                cur2.close()
                conn2.close()
            except Exception as e:
                try:
                    conn2.close()
                except Exception:
                    pass
                print('Erro ao obter usuario_id:', e)

        if usuario_id and curso_id:
            resp3 = client.post('/vinculacao_curso', data={'usuario': usuario_id, 'curso': curso_id}, follow_redirects=True)
            print('vinculacao_curso status:', resp3.status_code)
        else:
            print('vinculacao_curso pulada: usuario_id ou curso_id ausente')


def run_migracao_hash():
    # Migra senhas em texto puro para hash seguro (pbkdf2:sha256)
    conn = get_db_connection()
    if not conn:
        print('Falha ao conectar ao banco para migração de hash.')
        return
    atualizados = 0
    try:
        cur = conn.cursor()
        cur.execute("SELECT id_usuario, senha FROM usuario")
        rows = cur.fetchall()
        for uid, senha in rows:
            s = str(senha or '')
            # Se já parece hash conhecido, pula
            if s.startswith('pbkdf2:') or s.startswith('$argon2'):
                continue
            novo = generate_password_hash(s)
            cur.execute("UPDATE usuario SET senha = %s WHERE id_usuario = %s", (novo, uid))
            atualizados += 1
        conn.commit()
        cur.close()
        conn.close()
        print(f'Migração concluída. Senhas atualizadas: {atualizados}.')
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        print('Erro na migração de hash:', e)


def run_seed_admins():
    # Cria/atualiza administradores padrão informados pelo usuário
    admins = [
        {
            'nome': 'Larissa Alinny',
            'email': 'aalinny9@gmail.com',
            'cpf': '000.000.000-00',
            'senha': 'LA123@47'
        },
        {
            'nome': 'Arthur Madeira',
            'email': 'Arthurmad456@gmail.com',
            'cpf': '000.000.000-00',
            'senha': 'AM123@47'
        },
        {
            'nome': 'João Vitor Ferreira',
            'email': 'vitorjoao123z@gmail.com',
            'cpf': '000.000.000-00',
            'senha': 'JV123@47'
        },
        {
            'nome': 'Victor Hugo Freitas',
            'email': 'victorhugofreitas123@gmail.com',
            'cpf': '000.000.000-00',
            'senha': 'VH123@47'
        },
        {
            'nome': 'Renata Fagundes',
            'email': 'renata.facinpro@gmail.com',
            'cpf': '000.000.000-00',
            'senha': 'RF123@47'
        },
        {
            'nome': 'Livio Lucas',
            'email': 'liviool123@gmail.com',
            'cpf': '000.000.000-00',
            'senha': 'Fac@1470'
        },
    ]

    conn = get_db_connection()
    if not conn:
        print('Falha ao conectar ao banco para seed de administradores.')
        return
    try:
        for admin in admins:
            cur = conn.cursor(row_factory=dict_row)
            print(f"Verificando usuário: {admin['email']}")
            cur.execute("SELECT id_usuario, tipo FROM usuario WHERE email = %s", (admin['email'],))
            user = cur.fetchone()
            senha_hash = generate_password_hash(admin['senha'])
            if user:
                cur2 = conn.cursor()
                cur2.execute(
                    "UPDATE usuario SET nome = %s, cpf = %s, senha = %s, tipo = %s WHERE id_usuario = %s",
                    (admin['nome'], admin['cpf'], senha_hash, 'Funcionário', user['id_usuario'])
                )
                conn.commit()
                cur2.close()
                audit_log('seed_admin_updated', {'email': admin['email']})
                print(f"Admin atualizado: {admin['nome']} <{admin['email']}>")
            else:
                cur2 = conn.cursor()
                cur2.execute(
                    "INSERT INTO usuario (nome, email, cpf, senha, tipo, curso_usuario) VALUES (%s, %s, %s, %s, %s, %s)",
                    (admin['nome'], admin['email'], admin['cpf'], senha_hash, 'Funcionário', None)
                )
                conn.commit()
                cur2.close()
                audit_log('seed_admin_created', {'email': admin['email']})
                print(f"Admin criado: {admin['nome']} <{admin['email']}>")
            cur.close()
        conn.close()
        print('Seed de administradores concluído.')
    except Exception as e:
        try:
            conn.close()
        except Exception:
            pass
        print('Erro ao executar seed de administradores:', e)


@app.route('/relatorio/abnt')
def relatorio_abnt_doc():
    # Geração de Word (.doc) via HTML compatível
    html = render_template(
        'relatorio_abnt_doc.html',
        titulo='Relatório técnico de otimizações de desempenho front-end',
        projeto='Inprolib - Projeto TelaTest ADM',
        autor='Assistente IA',
        local='Brasil',
        ano='2025'
    )
    resp = make_response(html)
    resp.headers['Content-Type'] = 'application/msword; charset=utf-8'
    resp.headers['Content-Disposition'] = 'attachment; filename="Relatorio_ABNT_Inprolib.doc"'
    return resp

if __name__ == '__main__':
    # Executa a validação quando chamado com --validate; migração com --hash-migrate; caso contrário, sobe o servidor.
    if len(sys.argv) > 1:
        arg = sys.argv[1]
        if arg in ('--validate', 'validate'):
            run_validacao()
        elif arg in ('--hash-migrate', 'hash-migrate'):
            run_migracao_hash()
        elif arg in ('--seed-admins', 'seed-admins'):
            run_seed_admins()
        else:
            app.run(debug=True)
    else:
        app.run(debug=True)